import sys
import os
from pathlib import Path
import subprocess
from contextlib import asynccontextmanager
from fastapi import FastAPI, UploadFile, File, HTTPException, Response, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
from typing import List, Dict, Any, Optional, Union, Tuple
from dotenv import load_dotenv
import logging
import google.generativeai as genai
from io import BytesIO
import json
from datetime import datetime
import PyPDF2
import docx
import uvicorn
import random
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import SupabaseVectorStore
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.schema import Document
from langchain.prompts import PromptTemplate
from supabase import create_client, Client

# Add the current directory to Python path
current_dir = os.path.dirname(os.path.abspath(__file__))
if current_dir not in sys.path:
    sys.path.insert(0, current_dir)

# Local imports from the same directory
from stt import speech_to_text
from voice_assistant import clean_text_for_speech, generate_gemini_response
from youtube_utils import search_youtube_videos

# Import routers
from agents.personalization.router import router as personalization_router
from agents.explanation.router import router as explanation_router
from agents.explainer_agent import router as explainer_router
from agents.coding_agent import router as coding_router
from agents.code_fixer_agent import router as code_fixer_router
from fastapi.staticfiles import StaticFiles

# Import learning and analysis agents
from agents.enhanced_learning_agent import learning_agent, PersonalizedFeedback, SkillLevel
from agents.code_analysis_agent import advanced_analyzer, ComprehensiveCodeReview, CodeIssue

# Import new advanced agents
from agents.advanced_personalization_agent import advanced_personalization_agent, LearningInsight, UserCodingStyle, ProjectContext
from agents.code_refactoring_agent import code_refactoring_agent, RefactoringSuggestion, PerformanceOptimization, CodeQualityMetrics
from agents.security_bug_agent import security_bug_agent, SecurityIssue, BugReport
from agents.algorithm_explanation_agent import algorithm_explanation_agent, AlgorithmExplanation, CodeFlowStep, VisualRepresentation

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Note: Dynamic greeting generation is now handled by the PersonalizationAgent
# which generates contextually appropriate greetings based on user profile and interaction history

logger.info("Starting backend application initialization...")

# Verify environment variables
logger.info("Loading environment variables...")
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
# Replace DeepSeek variables with Gemini
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL_NAME = "gemini-2.0-flash" # Use the requested model name
YOUTUBE_API_KEY = os.getenv("YOUTUBE_API_KEY") # Load YouTube API Key

# Check all critical environment variables
missing_env_vars = []
if not SUPABASE_URL: missing_env_vars.append("SUPABASE_URL")
if not SUPABASE_KEY: missing_env_vars.append("SUPABASE_KEY")
if not SUPABASE_SERVICE_ROLE_KEY: missing_env_vars.append("SUPABASE_SERVICE_ROLE_KEY")
if not GEMINI_API_KEY: missing_env_vars.append("GEMINI_API_KEY")
if not YOUTUBE_API_KEY: missing_env_vars.append("YOUTUBE_API_KEY") # Check for YouTube API Key

if missing_env_vars:
    logger.error(f"Missing required environment variables: {', '.join(missing_env_vars)}")
    #exit(1) # Keep the app running for now even with missing keys, but log the error
else:
    logger.info("All required environment variables are present.")

@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("Application lifespan startup event triggered.")
    # Startup: Nothing to initialize for Gemini client
    yield
    # Shutdown
    try:
        # No need to close the Gemini client as it doesn't require explicit closing
        logger.info("Shutdown event: Gemini client doesn't require explicit closing")
    except Exception as e:
        logger.error(f"Error during shutdown: {e}")
    
    # Clean up Supabase client if needed
    logger.info("Application shutting down. Cleaning up resources.")

app = FastAPI(lifespan=lifespan)
logger.info("FastAPI app initialized.")

# Configure CORS
logger.info("Configuring CORS middleware...")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000", "http://localhost:3001", "http://127.0.0.1:3001", "http://localhost:5173", "http://127.0.0.1:5173"],  # Common development URLs
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
logger.info("CORS middleware configured.")

# Include the personalization router
logger.info("Including routers...")
app.include_router(personalization_router, prefix="/personalization", tags=["personalization"])

# Include the explanation router
app.include_router(explanation_router, prefix="/explanation", tags=["explanation"])

# Include the explainer router
app.include_router(explainer_router, prefix="/explainer", tags=["explainer"])

# Include the new coding agent router
app.include_router(coding_router, prefix="/coding", tags=["coding"])

# Include the new code fixer agent router
app.include_router(code_fixer_router, prefix="/code-fixer", tags=["code-fixer"])
logger.info("Routers included successfully.")

# Initialize Supabase client and create necessary tables
supabase: Client | None = None
logger.info("Initializing Supabase client...")
try:
    supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
    
    # Create service role client for admin operations (bypasses RLS)
    service_role_supabase = None
    if SUPABASE_SERVICE_ROLE_KEY:
        service_role_supabase = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)
        logger.info("Service role client initialized successfully")
    else:
        logger.warning("Service role key not found - profile operations may fail due to RLS")
    logger.info("Supabase client initialized successfully.")

    # Test connection by performing a simple query
    logger.info("Testing Supabase connection...")
    response = supabase.table("documents").select("id").limit(1).execute()
    if hasattr(response, 'error') and response.error:
        raise Exception(f"Supabase connection test failed: {response.error}")
    logger.info(f"Supabase connection test successful: {response}")

    # Ensure the chat_history table exists
    # Note: Supabase doesn't provide a direct API to create tables via the client library.
    # We can use the SQL API to create the table if it doesn't exist.
    create_table_query = """
    CREATE TABLE IF NOT EXISTS chat_history (
        id SERIAL PRIMARY KEY,
        filename TEXT NOT NULL UNIQUE,
        chat_data JSONB NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );
    """
    # Execute the query using the Supabase REST API's RPC endpoint (requires enabling `pgbouncer` or using direct SQL execution)
    # Alternatively, you can run this SQL in the Supabase SQL Editor manually for now.
    # For automation, we can use the `postgrest` RPC if available, but we'll log a message for manual setup.
    logger.info("Please ensure the 'chat_history' table exists. Run the following SQL in Supabase SQL Editor:")
    logger.info(create_table_query)

    # If you have the service role key, you can execute SQL directly (uncomment the following):
    # supabase.postgrest.rpc("execute_sql", {"query": create_table_query}).execute()
    # logger.info("chat_history table created or already exists.")

except Exception as e:
    logger.error(f"Failed to initialize Supabase client or create tables: {e}")
    logger.error("Application will exit due to database connection failure.")
    exit(1)  # Exit if Supabase connection fails

# Create a variable to hold the initialized Gemini model
initialized_gemini_model: genai.GenerativeModel | None = None
logger.info("Initializing Gemini API client...")
try:
    # Configure the Gemini API
    genai.configure(api_key=GEMINI_API_KEY)
    # Create an instance of the GenerativeModel
    initialized_gemini_model = genai.GenerativeModel(GEMINI_MODEL_NAME)
    logger.info(f"Gemini API client initialized successfully with model {GEMINI_MODEL_NAME}")
    # The client is configured globally, so we don't need a specific client object
    # We'll set a flag to indicate successful initialization
    # gemini_client = True # Remove or comment out this line
except Exception as e:
    logger.error(f"Failed to initialize Gemini API client or model: {e}")
    logger.error("Application will exit due to Gemini API client failure.")
    exit(1)  # Exit if Gemini client fails
    # gemini_client = None # Remove or comment out this line

# Attach the initialized Gemini model instance to app.state
logger.info("Attaching Gemini client to app state.")
# This makes it accessible via the get_gemini_client dependency
app.state.gemini_client = initialized_gemini_model

# Initialize embeddings
embeddings: GoogleGenerativeAIEmbeddings | None = None
logger.info("Initializing Google Generative AI embeddings...")
try:
    # Use the GEMINI_API_KEY directly here for the embeddings model
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GEMINI_API_KEY)
    logger.info("Google Generative AI embeddings initialized successfully.")
except Exception as e:
    logger.error(f"Failed to initialize Google Generative AI embeddings: {e}")
    logger.error("Application will exit due to embeddings initialization failure.")
    exit(1)  # Exit if embeddings fail

logger.info("Backend application initialization complete.")

# Initialize vector store
vector_store: SupabaseVectorStore | None = None
try:
    if supabase and embeddings:
        vector_store = SupabaseVectorStore(
            client=supabase,
            embedding=embeddings,
            table_name="documents",
            query_name="custom_vector_search"
        )
        logger.info("SupabaseVectorStore initialized successfully with query_name='custom_vector_search'.")
    else:
        logger.warning("Supabase or Embeddings not initialized. Vector store not created.")
except Exception as e:
    logger.error(f"Failed to initialize SupabaseVectorStore: {e}")
    vector_store = None

# Initialize text splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    length_function=len
)

# In-memory query history
query_history: List[Tuple[str, str]] = []

# In-memory storage for active game quizzes (simple, resets on server reload)
active_game_quizzes = {}

# Pydantic models
class QueryRequest(BaseModel):
    query: str
    use_source_only: bool = False
    user_context: Optional[Dict[str, Any]] = None  # Add user context for personalization
    user_id: Optional[str] = None  # Added user_id field for persistent context

class VoiceQueryRequest(BaseModel): # New model for voice queries
    text: str
    chat_history: List[Dict[str, Any]] = [] # Allow sending current chat history
    file_context: str | None = None # Allow sending file context if any
    feedback: str | None = None # ADDED feedback field

class SummarizeChatRequest(BaseModel): # New model for summarization requests
    chat_history: List[Dict[str, Any]]

class YouTubeSearchRequest(BaseModel): # New model for YouTube search requests
    query: str

class SaveChatRequest(BaseModel):
    filename: str
    chat_history: List[Dict[str, Any]]

class DeleteChatRequest(BaseModel):
    id: int

class UpdateChatFilenameRequest(BaseModel):
    id: int
    filename: str

class QuizRequest(BaseModel):
    topic: str
    difficulty: str = "medium"  # easy, medium, hard
    num_questions: int = 5
    question_type: str = "multiple_choice"  # multiple_choice, true_false, open_ended

class QuizQuestion(BaseModel):
    question: str
    options: List[str] = []
    correct_answer: str
    explanation: str
    points: int = 10

class QuizResponse(BaseModel):
    topic: str
    difficulty: str
    questions: List[QuizQuestion]
    total_points: int

class CodeGenerationRequest(BaseModel):
    prompt: str
    provider: str = "gemini"  # "gemini" or "azure" (for future expansion)

class CodeGenerationResponse(BaseModel):
    success: bool
    code: str
    provider: str = "gemini"
    error: Optional[str] = None

# Pydantic models
class CodeAnalysisRequest(BaseModel):
    code: str
    language: str = "javascript"
    user_id: Optional[str] = None
    user_level: str = "beginner"

class CodeAnalysisResponse(BaseModel):
    overall_score: float
    issues: List[dict]
    security_score: float
    performance_score: float
    educational_insights: List[str]
    improvement_roadmap: List[str]
    personalized_feedback: dict

class UserLearningProfile(BaseModel):
    user_id: str
    skill_level: str
    knowledge_areas: dict
    learning_progress: dict
    recommendations: List[str]
    learning_preferences: dict
    goals: List[str]
    weak_topics: List[str]
    interaction_types: dict
    topic_progress: dict
    achievements: List[str]

class SkillAssessmentRequest(BaseModel):
    user_id: str
    code_samples: List[str]
    language: str = "javascript"

class LearningInteractionRequest(BaseModel):
    user_id: str
    interaction_type: str
    topic: str
    code: Optional[str] = None
    complexity: str = "simple"

# Utility functions
def extract_pdf_text(file: BytesIO) -> str:
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = "".join(page.extract_text() or "" for page in pdf_reader.pages)
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        return ""

def extract_docx_text(file: BytesIO) -> str:
    try:
        doc = docx.Document(file)
        text = "\n".join([p.text for p in doc.paragraphs if p.text])
        return text
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        return ""

def process_file(file: bytes, file_type: str, filename: str) -> List[Document]:
    file_stream = BytesIO(file)
    raw_text = ""
    if file_type == "pdf":
        raw_text = extract_pdf_text(file_stream)
    elif file_type == "docx":
        raw_text = extract_docx_text(file_stream)
    else:
        logger.warning(f"Attempted to process unsupported file type: {file_type}")
        return []

    if not raw_text:
        logger.warning(f"No text extracted from {filename}")
        return []

    docs = [Document(page_content=raw_text, metadata={"filename": filename})]
    split_docs = text_splitter.split_documents(docs)
    logger.info(f"Split {filename} into {len(split_docs)} documents")
    return split_docs

# Updated OPTIONS handlers with explicit CORS headers and logging
@app.options("/upload")
async def options_upload():
    logger.info("Handling OPTIONS request for /upload")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "http://localhost:3000",
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
            "Access-Control-Allow-Credentials": "true"
        }
    )

@app.options("/query")
async def options_query():
    logger.info("Handling OPTIONS request for /query")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "http://localhost:3000",
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
            "Access-Control-Allow-Credentials": "true"
        }
    )

@app.options("/speech-to-text")
async def options_speech_to_text():
    logger.info("Handling OPTIONS request for /speech-to-text")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "http://localhost:3000",
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
            "Access-Control-Allow-Credentials": "true"
        }
    )

@app.options("/save-chat")
async def options_save_chat():
    logger.info("Handling OPTIONS request for /save-chat")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "http://localhost:3000",
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
            "Access-Control-Allow-Credentials": "true"
        }
    )

@app.options("/delete-chat")
async def options_delete_chat():
    logger.info("Handling OPTIONS request for /delete-chat")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "http://localhost:3000",
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
            "Access-Control-Allow-Credentials": "true"
        }
    )

@app.options("/update-chat-filename")
async def options_update_chat_filename():
    logger.info("Handling OPTIONS request for /update-chat-filename")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "http://localhost:3000",
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
            "Access-Control-Allow-Credentials": "true"
        }
    )

@app.options("/voice-query") # OPTIONS for new endpoint
async def options_voice_query():
    logger.info("Handling OPTIONS request for /voice-query")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "http://localhost:3000",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
            "Access-Control-Allow-Credentials": "true"
        }
    )

@app.options("/summarize-voice-chat") # OPTIONS for new summarization endpoint
async def options_summarize_voice_chat():
    logger.info("Handling OPTIONS request for /summarize-voice-chat")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "http://localhost:3000",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
            "Access-Control-Allow-Credentials": "true"
        }
    )

@app.options("/fetch-youtube-videos") # OPTIONS for new YouTube endpoint
async def options_fetch_youtube_videos():
    logger.info("Handling OPTIONS request for /fetch-youtube-videos")
    return Response(
            status_code=200,
                    headers={
                        "Access-Control-Allow-Origin": "http://localhost:3000",
                        "Access-Control-Allow-Methods": "POST, OPTIONS",
                        "Access-Control-Allow-Headers": "Content-Type, Authorization",
                        "Access-Control-Allow-Credentials": "true"
                    }
    )

@app.options("/generate-quiz")
async def options_generate_quiz():
    logger.info("Handling OPTIONS request for /generate-quiz")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "http://localhost:3000",
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
            "Access-Control-Allow-Credentials": "true"
        }
    )

@app.post("/stop-talking")
async def stop_talking():
    """
    Endpoint to signal that the text-to-speech on the client was stopped.
    This can be used for logging or state management on the backend.
    """
    logger.info("Received signal: client stopped text-to-speech.")
    # In a more advanced setup, this could trigger logic to
    # update user state, analytics, or stop a streaming response.
    return JSONResponse(status_code=200, content={"message": "Stop signal received and logged."})

@app.options("/stop-talking")
async def options_stop_talking():
    logger.info("Handling OPTIONS request for /stop-talking")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "http://localhost:3000",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
            "Access-Control-Allow-Credentials": "true"
        }
    )

@app.post("/save-chat")
async def save_chat(request: SaveChatRequest):
    logger.info(f"Received request to save chat with filename: {request.filename}")
    logger.debug(f"Chat history: {request.chat_history}")

    try:
        if supabase is None:
            logger.error("Supabase client not initialized. Cannot save chat.")
            raise HTTPException(status_code=500, detail="Database client not initialized. Please check server logs for more details.")

        try:
            chat_history_str = json.dumps(request.chat_history, default=str)
        except (TypeError, ValueError) as e:
            logger.error(f"Failed to serialize chat_history: {e}")
            raise HTTPException(status_code=400, detail=f"Invalid chat_history format: {str(e)}")

        logger.info(f"Checking if filename '{request.filename}' already exists in chat_history table")
        existing_chat_query = supabase.table("chat_history").select("id").eq("filename", request.filename)
        existing_chat_response = existing_chat_query.execute()

        if hasattr(existing_chat_response, 'error') and existing_chat_response.error:
            logger.error(f"Error checking existing chat filename '{request.filename}': {existing_chat_response.error}")
            raise HTTPException(status_code=500, detail=f"Database error during filename check: {existing_chat_response.error}")

        if existing_chat_response.data and len(existing_chat_response.data) > 0:
            chat_id_to_update = existing_chat_response.data[0]["id"]
            logger.info(f"Updating existing chat with filename: {request.filename} (ID: {chat_id_to_update})")

            updated_data = {
                "chat_data": chat_history_str,
                "filename": request.filename
            }
            response = supabase.table("chat_history").update(updated_data).eq("id", chat_id_to_update).execute()

            if hasattr(response, 'error') and response.error:
                logger.error(f"Failed to update chat with id {chat_id_to_update}: {response.error}")
                error_message = getattr(response.error, 'message', str(response.error))
                raise HTTPException(status_code=500, detail=f"Failed to update chat: Database operation failed: {error_message}")

            # If we reach here, the DB operation itself did not return an error in response.error.
            # Now, check response.count to understand what happened.
            # response.count can be an int (0, 1, ...) or None.
            if response.count is not None and response.count > 0:
                logger.info(f"Successfully updated chat with id: {chat_id_to_update}. Rows affected: {response.count}")
                return JSONResponse(status_code=200, content={"message": f"Chat updated successfully as '{request.filename}'"})
            elif response.count == 0:
                logger.info(f"Chat with id {chat_id_to_update} targeted, but 0 rows affected by update. This may mean the data was already identical, or the row was not found (e.g., deleted concurrently).")
                return JSONResponse(status_code=200, content={"message": f"Chat '{request.filename}' data is already current or no specific changes were applied."})
            else: # response.count is None
                # This is an ambiguous case from the client's perspective.
                # The DB operation didn't error, but we don't know how many rows were affected.
                # Assume the operation was accepted by the DB if no error was reported.
                logger.warning(f"Update for chat id {chat_id_to_update} (filename: '{request.filename}') completed without a database error, but the affected row count is unknown (None). Assuming the request was processed.")
                return JSONResponse(status_code=200, content={"message": f"Chat update for '{request.filename}' processed by the database; confirmation of changes pending."})

        else:
            logger.info(f"Inserting new chat with filename: {request.filename}")
            chat_data_to_insert = {
                "filename": request.filename,
                "chat_data": chat_history_str
            }
            response = supabase.table("chat_history").insert([chat_data_to_insert]).execute()

            if hasattr(response, 'error') and response.error:
                logger.error(f"Failed to insert chat with filename {request.filename}: {response.error}")
                raise HTTPException(status_code=500, detail=f"Failed to save new chat: Database operation failed: {response.error}")

            if response.data and len(response.data) > 0:
                logger.info(f"Successfully saved new chat with filename: {request.filename} (ID: {response.data[0].get('id')})")
                return JSONResponse(status_code=201, content={"message": f"Chat saved successfully as '{request.filename}'", "id": response.data[0].get('id')})
            else:
                logger.error("Failed to save new chat: Supabase insert returned no data.")
                raise HTTPException(status_code=500, detail="Failed to save new chat: Database operation failed.")

    except HTTPException as he:
        logger.error(f"HTTPException in save_chat: {he.detail}")
        raise he
    except Exception as e:
        logger.error(f"Unexpected error saving chat: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error saving chat: {str(e)}")

@app.post("/upload")
async def upload_files(files: List[UploadFile] = File(...)):
    processed_filenames = []
    failed_filenames = []
    try:
        if vector_store is None:
            logger.error("Vector store not initialized. Cannot process upload.")
            raise HTTPException(status_code=500, detail="Server not fully initialized. Vector store unavailable.")

        for file in files:
            file_content = await file.read()
            file_type = file.filename.split(".")[-1].lower()

            if file_type not in ["pdf", "docx"]:
                logger.warning(f"Rejected unsupported file type: {file.filename}")
                failed_filenames.append(f"{file.filename} (unsupported type)")
                continue

            try:
                split_docs = process_file(file_content, file_type, file.filename)

                if not split_docs:
                    logger.error(f"Failed to process {file.filename}: No documents generated.")
                    failed_filenames.append(f"{file.filename} (processing failed)")
                    continue

                logger.info(f"--- Sample Chunks for {file.filename} ---")
                for i, doc_chunk in enumerate(split_docs[:2]): # Log first 2 chunks
                    logger.info(f"Chunk {i+1} (first 100 chars): {doc_chunk.page_content[:100]}")
                    logger.info(f"Chunk {i+1} metadata: {doc_chunk.metadata}")
                logger.info(f"--- End Sample Chunks --- ({len(split_docs)} total chunks to add)")

                logger.info(f"--- Preparing to add {len(split_docs)} docs to Supabase for {file.filename} ---")
                for i, doc_to_add in enumerate(split_docs[:2]): # Log first 2 docs being sent
                    logger.info(f"Doc {i+1} Content (preview): {doc_to_add.page_content[:50]}...")
                    logger.info(f"Doc {i+1} METADATA being sent: {doc_to_add.metadata}") # THIS IS KEY
                logger.info(f"--- End Detailed Metadata Logging --- ")
                
                try:
                    vector_store.add_documents(split_docs)
                    logger.info(f"Successfully called vector_store.add_documents for {file.filename}")
                except Exception as vs_add_e:
                    logger.error(f"Error calling vector_store.add_documents for {file.filename}: {vs_add_e}", exc_info=True)
                    failed_filenames.append(f"{file.filename} (vector store add error)")
                    continue # Skip to next file if this one failed to add

                processed_filenames.append(file.filename)

            except Exception as file_e:
                logger.error(f"Error processing file {file.filename}: {file_e}")
                failed_filenames.append(f"{file.filename} (error: {file_e})")

        if not processed_filenames and failed_filenames:
            raise HTTPException(status_code=400, detail=f"Failed to process all files: {', '.join(failed_filenames)}")
        elif failed_filenames:
            return JSONResponse(status_code=207, content={"message": f"Processed {len(processed_filenames)} files successfully. Failed to process: {', '.join(failed_filenames)}"})
        else:
            return JSONResponse(status_code=200, content={"message": "Files processed successfully"})

    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Unexpected error during upload: {e}")
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during upload: {str(e)}")

RELATED_QUESTIONS_PROMPT = PromptTemplate.from_template(
    """
Based on this query and answer, generate exactly 3 related questions that would help the user explore this topic further. Return ONLY the questions, each on a new line starting with "- ".

Query: {query}
Answer: {answer}

- """
)

FALLBACK_RELATED_QUESTIONS = [
    "Can you provide more details on this topic?",
    "What are some examples related to this subject?",
    "How does this concept apply in real-world scenarios?"
]

async def process_single_query(query: str, use_source_only: bool = False) -> Dict[str, Any]:
    logger.info(f"Processing query: '{query}', use_source_only: {use_source_only}")

    # Initialize related_questions as empty list
    related_questions = []
    youtube_videos_list = []
    if vector_store is None:
        logger.error("Vector store not initialized. Cannot perform similarity search.")
        # Fetch YouTube videos even if vector store fails, as they might still be relevant
        if YOUTUBE_API_KEY:
            youtube_data = search_youtube_videos(query, max_results=4)
            if youtube_data and not youtube_data.get("error"):
                youtube_videos_list = youtube_data.get("videos", [])
            elif youtube_data.get("error"):
                 logger.error(f"Error fetching YouTube videos (vector store unavailable): {youtube_data.get('error')}")
        else:
            logger.warning("YOUTUBE_API_KEY not configured. Skipping video search.")

        return {
            "answer": "Error: The document vector store is not available. Please check server configuration.",
            "source_documents": [],
            "related_questions": FALLBACK_RELATED_QUESTIONS[:3],
            "youtube_videos": youtube_videos_list
        }

    serialized_source_documents = []
    source_documents = [] # Initialize to empty list
    source_documents_with_scores = []

    try:
        if embeddings:
            try:
                query_embedding = embeddings.embed_query(query)
                logger.info(f"Generated query embedding for '{query}' (first 5 dims): {query_embedding[:5]}, length: {len(query_embedding)}...")
            except Exception as qe_err:
                logger.error(f"Failed to generate query embedding: {qe_err}", exc_info=True)
                # Fallback or re-raise, depending on desired behavior. For now, log and continue, retrieval will likely fail.
                # This might be a good place to return an error specific to embedding failure.

        logger.info(f"Attempting basic similarity search for query: '{query}' with k=5 using asimilarity_search")
        # Using asimilarity_search (does not return scores directly in the same way)
        source_documents = await vector_store.asimilarity_search(query, k=5) # MODIFIED LINE
        
        logger.info(f"Retrieved {len(source_documents)} raw results from vector store (using asimilarity_search) for '{query}'.") # MODIFIED LOG
        
        if not source_documents:
            logger.warning(f"No documents returned by asimilarity_search for query: '{query}'. Attempting direct RPC call to match_documents.")
            if embeddings and query_embedding: # Ensure we have the embedding
                try:
                    rpc_params = {
                        "query_embedding": query_embedding,
                        "match_count": 5,
                        "filter": {}
                    }
                    logger.info(f"Attempting direct RPC call to 'custom_vector_search' with params: query_embedding (first 5 dims): {query_embedding[:5]}, match_count: 5, filter: {{}}")
                    
                    if supabase:
                        direct_match_response = supabase.rpc("custom_vector_search", rpc_params).execute()
                        logger.info(f"Direct RPC call response data: {direct_match_response.data}")
                        if hasattr(direct_match_response, 'error') and direct_match_response.error:
                            logger.error(f"Direct RPC call error: {direct_match_response.error}")
                        if direct_match_response.data:
                            logger.info(f"Direct RPC call to 'custom_vector_search' FOUND {len(direct_match_response.data)} documents.")
                            # For debugging, let's populate source_documents from this direct call if it succeeds
                            # This is a temporary measure for testing the RAG pipeline end-to-end if direct call works
                            # We'd need to map the structure if it's different from Langchain's Document
                            # For now, just logging is the primary goal. If successful, we can adapt.
                            # source_documents = [Document(page_content=item['content'], metadata=item['metadata']) for item in direct_match_response.data]
                            # logger.info(f"Populated source_documents with {len(source_documents)} items from direct RPC call for further processing.")
                        else:
                            logger.warning("Direct RPC call to 'custom_vector_search' also returned no documents or an error occurred.")
                    else:
                        logger.error("Supabase client is None, cannot make direct RPC call.")

                except Exception as rpc_err:
                    logger.error(f"Exception during direct RPC call to 'custom_vector_search': {rpc_err}", exc_info=True)
            else:
                logger.warning("Skipping direct RPC call because embeddings or query_embedding is not available.")
        
        if not source_documents: # Re-check after the RPC attempt (if we were to populate source_documents from RPC)
             logger.warning(f"Still no documents after asimilarity_search and potential direct RPC call for query: '{query}'")
        else:
            for i, doc in enumerate(source_documents): # MODIFIED LOOP
                logger.info(f"  Result {i+1}: Metadata={doc.metadata}, Content (first 50)='{doc.page_content[:50]}...'") # MODIFIED LOG

        # Since asimilarity_search returns documents directly, no score-based filtering needed here for this test
        # source_documents = [doc for doc, score in source_documents_with_scores] # Original line, commented out for now
        logger.info(f"{len(source_documents)} documents after basic retrieval (asimilarity_search).") # MODIFIED LOG

        serialized_source_documents = [
            {"page_content": doc.page_content, "metadata": doc.metadata} for doc in source_documents
        ]

    except Exception as e:
        logger.error(f"Error during document retrieval or embedding for query '{query}': {e}", exc_info=True)
        # source_documents will remain empty or partially filled if error occurred mid-process
        # The subsequent check `if use_source_only and not source_documents:` will handle this scenario.

    # This check is CRITICAL and is what you are seeing
    logger.debug(f"DEBUG in process_single_query: use_source_only={use_source_only}, source_documents_count={len(source_documents)}")
    if use_source_only and not source_documents:
        logger.info("Source-only mode active and no documents (source_documents list is empty). Returning specific message.")
        return {
            "answer": "No relevant documents were found in the knowledge base to answer your query based on the available sources.",
            "source_documents": [],
            "related_questions": FALLBACK_RELATED_QUESTIONS[:3], # Fallback questions
            "youtube_videos": [] 
        }

    context = "\n\n".join([doc.page_content for doc in source_documents])
    limited_history = query_history[-5:]
    history_str = "\n".join([f"User: {q}\nAssistant: {a}" for q, a in limited_history])

    # Adjust prompt based on use_source_only
    if use_source_only:
        prompt_template_str = """Based solely on the following context, answer the user's query. 
If the context does not contain the information to answer the query, state that the information is not available in the provided documents. Do not use any external knowledge.

Context:
{context}

Query: {query}

Answer:"""
    else:
        prompt_template_str = """You are a helpful assistant. Use the following context and chat history to answer the user's query. 
If the context is empty or not relevant, answer the query using your general knowledge.

Context:
{context}

Chat History:
{history}

Query: {query}

Answer:"""
    
    prompt = prompt_template_str.format(context=context, history=history_str, query=query)

    # Create system prompt text instead of using SystemMessage class
    system_prompt = "You are a helpful assistant."
    if use_source_only:
        system_prompt = "You are an assistant that answers questions based ONLY on the provided context. If the answer is not in the context, say so explicitly."

    raw_answer = ""
    try:
        if initialized_gemini_model is None:
            # Log error but try to continue for videos
            logger.error("Gemini client not initialized. Cannot generate LLM answer.")
            raw_answer = "Error: The language model is not available to generate an answer."
        else:
            logger.info(f"Attempting to query Gemini with prompt for query: {query}")
            
            # Send the prompt to Gemini
            response = initialized_gemini_model.generate_content(
                f"{system_prompt}\n\n{prompt}",
                generation_config={
                    "max_output_tokens": 12000,
                    "temperature": 0.8,
                    "top_p": 0.1
                }
            )
            
            # Get the response text
            raw_answer = response.text
            logger.info(f"Raw answer from Gemini: '{raw_answer[:200]}...' ({len(raw_answer)} chars)")
            
            if not raw_answer.strip():
                raw_answer = "Could not retrieve an answer from the language model."
                logger.warning("Gemini returned an empty answer.")
    except Exception as e:
        logger.error(f"Error querying Gemini for answer: {e}", exc_info=True)
        raw_answer = f"Error: Failed to get answer from the language model: {str(e)}"

    # Modify related questions generation if in source-only mode and no answer was found from sources
    if use_source_only and raw_answer.startswith("Error: Failed to get answer") or (not raw_answer.strip() and not source_documents):
        related_questions = ["Try rephrasing your query.", "Upload more documents that might cover this topic.", "Ask a question about a different topic based on the uploaded documents."]
    elif use_source_only and (raw_answer.strip() and not raw_answer.startswith("Error:")):
        # If source-only and got an answer, generate related questions normally or adapt the prompt
        try:
            if initialized_gemini_model is None:
                logger.warning("Gemini client not initialized. Skipping related questions.")
                related_questions = FALLBACK_RELATED_QUESTIONS[:3]
            else:
                logger.info(f"Attempting to generate related questions for query (source-only answer): {query}")
                related_questions_input = {"query": query, "answer": raw_answer}
                # Potentially a different prompt for related questions in source-only mode if needed
                related_questions_prompt_text = RELATED_QUESTIONS_PROMPT.format(**related_questions_input) 
                
                # Send the prompt for related questions
                logger.info("Calling Gemini for related questions (source-only answer)...")
                response_related = initialized_gemini_model.generate_content(
                    related_questions_prompt_text,
                    generation_config={
                        "max_output_tokens": 200,
                        "temperature": 0.8,
                        "top_p": 0.9
                    }
                )
                
                related_questions_text = response_related.text
                logger.info(f"Raw related questions response from Gemini (source-only answer): {related_questions_text}")
                
                if related_questions_text:
                    lines = related_questions_text.strip().split("\n")
                    for line in lines:
                        line = line.strip()
                        # Skip empty lines and introductory text
                        if not line or line.lower().startswith(("here are", "based on", "suggested", "questions:")):
                            continue
                        
                        # Extract question from different formats
                        question = ""
                        if line.startswith(("- ", "* ")):
                            question = line[2:].strip()
                        elif len(line) > 2 and line[0].isdigit() and line[1] in (".", ")", " "):
                            # Handle numbered lists like "1. Question"
                            idx = 0
                            while idx < len(line) and (line[idx].isdigit() or line[idx] in (".", ")", " ")):
                                idx += 1
                            question = line[idx:].strip()
                        elif line and not line.lower().startswith(("here", "based", "suggested")):
                            question = line.strip()
                        
                        # Add question if valid and unique
                        if question and len(related_questions) < 3:
                            # Check for duplicates (case-insensitive)
                            if not any(question.lower() == existing.lower() for existing in related_questions):
                                related_questions.append(question)
                        
                        if len(related_questions) >= 3:
                            break
                fallback_index = 0
                while len(related_questions) < 3 and fallback_index < len(FALLBACK_RELATED_QUESTIONS):
                    fb_q = FALLBACK_RELATED_QUESTIONS[fallback_index]
                    if all(fb_q.lower() != existing.lower() for existing in related_questions):
                        related_questions.append(fb_q)
                    fallback_index += 1
                logger.info(f"Generated related questions (source-only answer): {related_questions}")
        except Exception as e:
            logger.error(f"Error generating related questions with Gemini (source-only answer): {e}", exc_info=True)
            related_questions = FALLBACK_RELATED_QUESTIONS[:3]
            logger.info(f"Using fallback related questions due to error (source-only answer): {related_questions}")
    else: # Original related questions logic for non-source-only mode
        try:
            if initialized_gemini_model is None:
                logger.warning("Gemini client not initialized. Skipping related questions.")
                related_questions = FALLBACK_RELATED_QUESTIONS[:3]
            else:
                logger.info(f"Attempting to generate related questions for query: {query}")
                related_questions_input = {"query": query, "answer": raw_answer}
                related_questions_prompt_text = RELATED_QUESTIONS_PROMPT.format(**related_questions_input)
                
                # Send the prompt for related questions
                logger.info("Calling Gemini for related questions...")
                response_related = initialized_gemini_model.generate_content(
                    related_questions_prompt_text,
                    generation_config={
                        "max_output_tokens": 200,
                        "temperature": 0.8,
                        "top_p": 0.9
                    }
                )
                
                related_questions_text = response_related.text
            
                logger.info(f"Raw related questions response from Gemini: {related_questions_text}")
                if related_questions_text:
                    lines = related_questions_text.strip().split("\n")
                    for line in lines:
                        line = line.strip()
                        # Skip empty lines and introductory text
                        if not line or line.lower().startswith(("here are", "based on", "suggested", "questions:")):
                            continue
                        
                        # Extract question from different formats
                        question = ""
                        if line.startswith(("- ", "* ")):
                            question = line[2:].strip()
                        elif len(line) > 2 and line[0].isdigit() and line[1] in (".", ")", " "):
                            # Handle numbered lists like "1. Question"
                            idx = 0
                            while idx < len(line) and (line[idx].isdigit() or line[idx] in (".", ")", " ")):
                                idx += 1
                            question = line[idx:].strip()
                        elif line and not line.lower().startswith(("here", "based", "suggested")):
                            question = line.strip()
                        
                        # Add question if valid and unique
                        if question and len(related_questions) < 3:
                            # Check for duplicates (case-insensitive)
                            if not any(question.lower() == existing.lower() for existing in related_questions):
                                related_questions.append(question)
                        
                        if len(related_questions) >= 3:
                            break
            
            fallback_index = 0
            while len(related_questions) < 3 and fallback_index < len(FALLBACK_RELATED_QUESTIONS):
                fb_q = FALLBACK_RELATED_QUESTIONS[fallback_index]
                if all(fb_q.lower() != existing.lower() for existing in related_questions):
                    related_questions.append(fb_q)
                fallback_index += 1
            logger.info(f"Generated related questions: {related_questions}")

        except Exception as e:
            logger.error(f"Error generating related questions with Gemini: {e}", exc_info=True)
            related_questions = FALLBACK_RELATED_QUESTIONS[:3]
            logger.info(f"Using fallback related questions due to error: {related_questions}")
    
    related_questions = related_questions[:3] # Ensure only 3
    while len(related_questions) < 3: # Ensure exactly 3, even if duplicates were avoided
        related_questions.append("Explore this topic further.")


    # Fetch YouTube videos
    if YOUTUBE_API_KEY:
        logger.info(f"Attempting to fetch YouTube videos for query: {query}")
        youtube_data = search_youtube_videos(query, max_results=4)
        if youtube_data and not youtube_data.get("error"):
            youtube_videos_list = youtube_data.get("videos", [])
            logger.info(f"Successfully fetched {len(youtube_videos_list)} YouTube videos.")
        elif youtube_data.get("error"):
            logger.error(f"Error fetching YouTube videos for query '{query}': {youtube_data.get('error')}")
            # youtube_videos_list remains empty as initialized
    else:
        logger.warning("YOUTUBE_API_KEY not configured. Skipping video search.")


    query_history.append((query, raw_answer)) # Ensure query_history is defined and accessible

    logger.info(f"Returning final processed response for query: {query} - Answer: '{raw_answer[:100]}...' Videos: {len(youtube_videos_list)}")
    return {
        "answer": raw_answer,
        "source_documents": serialized_source_documents,
        "related_questions": related_questions,
        "youtube_videos": youtube_videos_list # Add fetched videos to the response
    }

@app.post("/query")
async def query_documents_endpoint(request: QueryRequest):
    try:
        user_id = request.user_id or "guest"
        logger.info(f"Received query: '{request.query}', from user: {user_id}, with user context: {bool(request.user_context)}")
        
        # We'll track the query and response together at the end of the function
        # to avoid duplicate tracking
        
        # Extract user ID from user context, or generate a fallback
        user_id = "guest_user"
        if request.user_context and request.user_context.get('user', {}).get('name'):
            user_email = request.user_context.get('user', {}).get('name', '')
            if '@' in user_email:
                # Use email as user identifier
                user_id = user_email
            else:
                # Use name as identifier
                user_id = f"user_{user_email}"
        else:
            # Fallback for guest users
            user_id = "guest_" + str(hash(request.query))[:8]
        
        # Enhanced personalization with user context
        if request.user_context:
            logger.info(f"Processing with user context: {request.user_context}")
            
            # Create personalized prompt based on user context
            user_profile = request.user_context.get('user', {})
            session_info = request.user_context.get('session', {})
            context_info = request.user_context.get('context', {})
            
            # Extract user interests/topics for personalization
            user_interests = user_profile.get('interests', [])
            user_topics = user_profile.get('topicsOfInterest', [])
            user_skills = user_profile.get('currentSkills', [])
            weak_topics = user_profile.get('weakTopics', [])
            
            # Tailor the query based on user preferences
            personalized_instructions = []
            
            learning_style = user_profile.get('learningStyle', 'unknown')
            skill_level = user_profile.get('skillLevel', 'beginner')
            preferred_difficulty = user_profile.get('preferredDifficulty', 'medium')
            
            # Add emojis based on learning style
            if learning_style == 'visual':
                personalized_instructions.append("📊 Include examples, diagrams, and visual explanations where possible.")
            elif learning_style == 'auditory':
                personalized_instructions.append("🎧 Provide step-by-step verbal explanations and use clear, spoken-friendly language.")
            elif learning_style == 'kinesthetic':
                personalized_instructions.append("🛠️ Focus on hands-on examples and practical exercises.")
            
            if skill_level == 'beginner':
                personalized_instructions.append("Explain concepts from the basics, avoid jargon, and provide simple examples.")
            elif skill_level == 'intermediate':
                personalized_instructions.append("Provide moderate detail with some advanced concepts and practical applications.")
            elif skill_level == 'advanced':
                personalized_instructions.append("Focus on advanced concepts, best practices, and optimization techniques.")
            
            if preferred_difficulty == 'easy':
                personalized_instructions.append("Keep explanations simple and easy to understand.")
            elif preferred_difficulty == 'hard':
                personalized_instructions.append("Provide comprehensive, detailed explanations with advanced insights.")
            
            # Include recent topics for context
            recent_topics = user_profile.get('recentTopics', [])
            if recent_topics:
                personalized_instructions.append(f"Consider the user's recent learning topics: {', '.join(recent_topics[:3])}")
            
            # Add session context
            if session_info.get('concepts'):
                personalized_instructions.append(f"Build upon previously discussed concepts: {', '.join(session_info['concepts'][:3])}")
            
            # Create enhanced query with personalization
            if personalized_instructions:
                enhanced_query = f"{request.query}\n\nPersonalization Instructions: {' '.join(personalized_instructions)}"
                
                # Add context about user's knowledge and interests
                if user_interests:
                    enhanced_query += f"\n\nUser has shown interest in: {', '.join(user_interests[:5])}"
                if user_skills:
                    enhanced_query += f"\n\nUser's current skills include: {', '.join(user_skills[:3])}"
                if weak_topics:
                    enhanced_query += f"\n\nUser wants to improve in: {', '.join(weak_topics[:3])}"
                
                logger.info(f"Enhanced query with personalization: {enhanced_query[:200]}...")
            else:
                enhanced_query = request.query
        else:
            enhanced_query = request.query
        
        try:
            # Extract user context information
            user_name = "there"  # Default fallback
            learning_style = "beginner"
            user_context_data = {}
            
            if request.user_context:
                user_profile = request.user_context.get('user', {})
                user_name = user_profile.get('name', user_name)
                learning_style = user_profile.get('learningStyle', learning_style)
                skill_level = user_profile.get('skillLevel', 'beginner')
                user_context_data = request.user_context
                
                logger.info(f"User context found - Name: {user_name}, Learning Style: {learning_style}, Skill Level: {skill_level}")
            
            # Check if this is a simple greeting or "about me" query
            query_lower = request.query.lower().strip()
            
            # Handle "about me" and profile-related queries directly with user context
            profile_query_phrases = [
                "about me", "do you know me", "who am i", "my information", "what do you know about me",
                "my goals", "what are my goals", "my learning goals", "my objectives", 
                "my skills", "what are my skills", "my interests", "my preferences",
                "my background", "my experience", "my education", "my profile",
                "tell me about my", "what's my", "show me my", "in your memory", "what do you remember",
                "our conversation", "what we discussed", "conversation history", "what did we talk about",
                "remember", "yesterday", "earlier", "before", "previously", "last time"
            ]
            
            if any(phrase in query_lower for phrase in profile_query_phrases):
                if user_name and user_name != "there":
                    personalized_about = f"Yes! I have quite a bit in my memory about you, {user_name}. "
                    
                    # Check for conversation history in personalization system
                    conversation_history = []
                    try:
                        if request.user_id:
                            from agents.personalization.user_context import get_user_context
                            user_context = get_user_context(request.user_id)
                            conversation_history = user_context.context.get('conversationHistory', [])
                            logger.info(f"Found {len(conversation_history)} conversations in history for {request.user_id}")
                    except Exception as e:
                        logger.warning(f"Could not retrieve conversation history: {e}")
                    
                    # Add conversation history if available
                    if conversation_history:
                        personalized_about += f"We've had {len(conversation_history)} conversations recently. "
                        
                        # Show recent topics discussed
                        recent_topics = []
                        for conv in conversation_history[-5:]:  # Last 5 conversations
                            if conv.get('topic'):
                                recent_topics.append(conv['topic'])
                        
                        if recent_topics:
                            unique_topics = list(set(recent_topics))  # Remove duplicates
                            personalized_about += f"We've discussed topics like: {', '.join(unique_topics[:3])}. "
                        
                        # Show last conversation if available
                        if conversation_history:
                            last_conv = conversation_history[-1]
                            if last_conv.get('query'):
                                personalized_about += f"In our last conversation, you asked about: '{last_conv['query'][:50]}...' "
                    
                    # Add profile details if available
                    user_profile = user_context_data.get('user', {})
                    details = []
                    
                    # Provide detailed context if available
                    if user_profile.get('skillLevel') == 'expert':
                        details.append("You're an expert in your field!")
                    
                    if user_profile.get('age'):
                        details.append(f"You're {user_profile['age']} years old")
                    
                    if user_profile.get('education'):
                        details.append(f"your education level is {user_profile['education']}")
                    
                    if user_profile.get('occupation'):
                        details.append(f"using your knowledge in your current role as a(n) {user_profile['occupation']}")
                    
                    if user_profile.get('currentSkills'):
                        skills = user_profile['currentSkills'][:3]  # Show top 3 skills
                        if skills:
                            details.append(f"you have skills in {', '.join(skills)}")
                    
                    if user_profile.get('learningGoals'):
                        goals = user_profile['learningGoals'][:2]  # Show top 2 goals
                        if goals:
                            details.append(f"your learning goals include {', '.join(goals)}")
                    
                    if user_profile.get('interests'):
                        interests = user_profile['interests'][:3]  # Show top 3 interests
                        if interests:
                            details.append(f"you're interested in {', '.join(interests)}")
                    
                    if user_profile.get('timeAvailable'):
                        details.append(f"you have {user_profile['timeAvailable']} hours available for learning per week")
                    
                    if details:
                        personalized_about += ". ".join(details) + ". "
                    
                    if learning_style and learning_style != "unknown":
                        personalized_about += f"You prefer {learning_style} learning, "
                    if skill_level:
                        personalized_about += f"and are at a {skill_level} skill level. "
                    
                    personalized_about += "All this information helps me personalize your learning experience and remember our conversations!"
                else:
                    personalized_about = "I can see you're using the system, but I don't have specific details about your profile yet. You can set up your learning preferences to get a more personalized experience!"
                
                response_data = {
                    "answer": personalized_about,
                    "source_documents": [],
                    "related_questions": [
                        "How can I update my learning preferences?",
                        "What learning styles do you support?",
                        "How does personalization work?"
                    ]
                }
                return JSONResponse(status_code=200, content=response_data)
            
            # Handle simple greetings with personalization (check for actual greetings, not substrings)
            greeting_phrases = ["hello", "hi", "hey", "good morning", "good afternoon", "good evening"]
            is_greeting = any(query_lower.strip().startswith(phrase) or query_lower.strip() == phrase for phrase in greeting_phrases)
            if is_greeting:
                personalized_greeting = f"Hello {user_name}! I'm your AI learning assistant. How can I help you with your studies today?"
                if user_name != "there" and user_name != "Guest User":
                    personalized_greeting = f"Hi {user_name}! Great to see you back. What would you like to learn about today?"
                elif user_name == "Guest User":
                    personalized_greeting = f"Welcome, {user_name}! I'm here to help you learn. What topics interest you?"
                
                response_data = {
                    "answer": personalized_greeting,
                    "source_documents": [],
                    "related_questions": [
                        "What topics would you like to learn about today?",
                        "Do you have any specific questions about a subject?",
                        "Would you like to see examples of what I can help with?"
                    ]
                }
                return JSONResponse(status_code=200, content=response_data)
            
            # First, route through personalization agent to get personalized instructions
            personalization_data = None
            try:
                from agents.personalization_agent import PersonalizationAgent
                
                # Ensure user_id is valid
                if not user_id or user_id == "guest_user":
                    # For guest users, use a default profile
                    logger.info(f"Using guest profile for user_id: {user_id}")
                    personalization_data = {
                        "query_type": "educational",
                        "level": "beginner",
                        "learning_style": ["visual", "textual"],
                        "tailored_query": enhanced_query
                    }
                else:
                    logger.info(f"Creating PersonalizationAgent for user_id: {user_id}")
                    agent = PersonalizationAgent(user_id)
                    
                    # Process the query and get personalization data
                    personalization_data = agent.process_query(enhanced_query)
                    logger.info(f"Personalization data for query: {personalization_data}")
                    
            except Exception as personalization_error:
                logger.warning(f"Error using personalization agent: {personalization_error}")
                # Fallback to basic personalization
                personalization_data = {
                    "query_type": "educational",
                    "level": skill_level if 'skill_level' in locals() else "beginner",
                    "learning_style": [learning_style] if 'learning_style' in locals() else ["visual"],
                    "tailored_query": enhanced_query
                }
            
            # Check if this is a greeting, non-educational, or profile query from the agent
            if personalization_data.get("query_type", "") == "greeting":
                # Use our enhanced greeting instead of the agent's basic one
                personalized_greeting = f"Hello {user_name}! I'm your AI learning assistant. How can I help you with your studies today?"
                if user_name != "there" and user_name != "Guest User":
                    personalized_greeting = f"Hi {user_name}! Great to see you back. What would you like to learn about today?"
                elif user_name == "Guest User":
                    personalized_greeting = f"Welcome, {user_name}! I'm here to help you learn. What topics interest you?"
                
                response_data = {
                    "answer": personalized_greeting,
                    "source_documents": [],
                    "related_questions": [
                        "What topics would you like to learn about today?",
                        "Do you have any specific questions about a subject?",
                        "Would you like to see examples of what I can help with?"
                    ]
                }
                return JSONResponse(status_code=200, content=response_data)
            
            # Handle profile/memory queries with detailed user information
            elif personalization_data.get("query_type", "") == "profile_query":
                # Create detailed profile response
                profile_response = personalization_data.get("response", "")
                if not profile_response:
                    # Generate detailed profile response if agent didn't provide one
                    user_profile = user_context_data.get('user', {})
                    profile_details = []
                    
                    if user_profile.get('name') and user_profile['name'] != "there":
                        profile_details.append(f"Your name is {user_profile['name']}")
                    
                    if user_profile.get('age'):
                        profile_details.append(f"you're {user_profile['age']} years old")
                    
                    if user_profile.get('education'):
                        profile_details.append(f"your education level is {user_profile['education']}")
                    
                    if user_profile.get('occupation'):
                        profile_details.append(f"you work as {user_profile['occupation']}")
                    
                    if user_profile.get('currentSkills'):
                        skills = user_profile['currentSkills'][:5]  # Show top 5 skills
                        if skills:
                            profile_details.append(f"you have skills in {', '.join(skills)}")
                    
                    # Add current goal (primary goal)
                    if user_profile.get('currentGoal'):
                        profile_details.append(f"your main goal is to {user_profile['currentGoal']}")
                    
                    # Add primary reason for learning
                    if user_profile.get('primaryReason'):
                        profile_details.append(f"you're learning for {user_profile['primaryReason']}")
                    
                    if user_profile.get('learningGoals'):
                        goals = user_profile['learningGoals'][:3]  # Show top 3 goals
                        if goals:
                            profile_details.append(f"your learning goals include {', '.join(goals)}")
                    
                    # Check multiple field names for interests/topics
                    interests = user_profile.get('interests') or user_profile.get('topicsOfInterest') or []
                    if interests:
                        interests_list = interests[:4]  # Show top 4 interests
                        profile_details.append(f"you're interested in {', '.join(interests_list)}")
                    
                    # Add experience level if different from skill level
                    if user_profile.get('experienceLevel') and user_profile.get('experienceLevel') != user_profile.get('skillLevel'):
                        profile_details.append(f"your experience level is {user_profile['experienceLevel']}")
                    
                    if user_profile.get('timeAvailable'):
                        profile_details.append(f"you have {user_profile['timeAvailable']} hours available for learning per week")
                    
                    if user_profile.get('motivation') and user_profile['motivation'].strip():
                        profile_details.append(f"your motivation is: {user_profile['motivation']}")
                    
                    if profile_details:
                        profile_response = f"Yes, I have quite a bit of information about you in my memory! {'. '.join(profile_details).capitalize()}. "
                    else:
                        profile_response = "I don't have much detailed information about you in my memory yet. "
                    
                    if learning_style and learning_style != "unknown":
                        profile_response += f"You prefer {learning_style} learning, "
                    if skill_level:
                        profile_response += f"and you're at a {skill_level} skill level. "
                    
                    profile_response += "I use all this information to personalize your learning experience and provide more relevant, tailored responses!"
                
                response_data = {
                    "answer": profile_response,
                    "source_documents": [],
                    "related_questions": [
                        "How can I update my learning preferences?",
                        "What learning goals should I set?",
                        "How does this personalization help my learning?"
                    ]
                }
                return JSONResponse(status_code=200, content=response_data)
            
            # For educational queries, use the personalization data to guide RAG
            tailored_query = personalization_data.get("tailored_query", enhanced_query)
            logger.info(f"Using tailored query from personalization agent: '{tailored_query[:100]}...'")
            
            # Pass the tailored query and use_source_only flag to process_single_query
            response_data = await process_single_query(tailored_query, request.use_source_only)
            
            # **CRITICAL**: Use the personalized greeting from the PersonalizationAgent
            personalized_greeting = personalization_data.get("personalized_greeting", "")
            
            if personalized_greeting and response_data.get("answer"):
                # Clean the personalized greeting to avoid including full query
                original_answer = response_data["answer"]
                
                # Extract just the greeting part, not the full instructions
                greeting_parts = personalized_greeting.split('\n')
                clean_greeting = greeting_parts[0]  # Take only the first line as greeting
                
                # Clean up any instruction text that might have leaked
                if "Personalization Instructions:" in clean_greeting:
                    clean_greeting = clean_greeting.split("Personalization Instructions:")[0].strip()
                
                response_data["answer"] = clean_greeting + "\n\n" + original_answer
                logger.info(f"Added personalized greeting to response: {clean_greeting[:50]}...")
            elif response_data.get("answer") and request.user_context:
                # Fallback personalization if no greeting from agent
                original_answer = response_data["answer"]
                
                # Try to use the new personalization system for more advanced adaptation
                try:
                    if request.user_id:
                        from agents.personalization import adapt_response_for_user
                        adapted_response = adapt_response_for_user(request.user_id, original_answer, request.query)
                        response_data["answer"] = adapted_response
                        logger.info(f"Response adapted using personalization system for user: {request.user_id}")
                    else:
                        # Create a simple friendly greeting if no personalization available
                        display_name = user_name.split('@')[0] if '@' in user_name and user_name != "there" else "there"
                        if display_name != "there":
                            simple_greeting = f"Hi {display_name}! "
                        else:
                            simple_greeting = "Hi there! "
                        
                        response_data["answer"] = simple_greeting + original_answer
                        
                except Exception as e:
                    logger.error(f"Error adapting response with personalization system: {e}")
                    # Fall back to original response
                    response_data["answer"] = original_answer
            
            # Note: tailored_instruction is used internally only, not exposed to user
            
            # Add user context info to response for debugging
            if request.user_context:
                response_data["debug_context"] = {
                    "learning_style": user_profile.get('learningStyle'),
                    "skill_level": user_profile.get('skillLevel'),
                    "personalized": True
                }
            
            # Track the response in personalization system if user_id is provided
            if request.user_id and response_data.get("answer"):
                try:
                    from agents.personalization.user_context import get_user_context
                    user_context = get_user_context(request.user_id)
                    user_context.update_from_query(request.query, response_data["answer"])
                    logger.info(f"Query and response tracked in personalization system for user: {request.user_id}")
                except Exception as e:
                    logger.warning(f"Error tracking response in personalization system: {e}")
            
            return JSONResponse(status_code=200, content=response_data)
            
        except ImportError:
            logger.warning("PersonalizationAgent not available, falling back to direct query processing")
        
        response_data = await process_single_query(enhanced_query, request.use_source_only)
        return JSONResponse(status_code=200, content=response_data)
            
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error processing query '{request.query}': {str(e)}")
        return JSONResponse(
            status_code=500,
            content={
                "answer": f"An error occurred while processing your query: {str(e)}",
                "source_documents": [],
                "related_questions": FALLBACK_RELATED_QUESTIONS[:3]
            }
        )

@app.post("/speech-to-text")
async def process_speech_to_text(audio_file: UploadFile = File(...)):
    try:
        # Read the uploaded file content
        audio_content = await audio_file.read()
        
        # Process the audio content
        result = speech_to_text(audio_content)
        
        if result.get("error"):
            raise HTTPException(status_code=400, detail=result.get("message"))
            
        return {"text": result.get("text")}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Voice assistant routes
class VoiceQuery(BaseModel):
    text: str
    chat_history: Optional[List[Dict[str, Any]]] = None

@app.post("/voice-query")
async def handle_voice_query(request: Request):
    try:
        # Get the JSON body from the request
        body = await request.json()
        user_text = body.get("text")
        chat_history = body.get("chat_history", [])
        
        if not user_text:
            raise HTTPException(status_code=400, detail="No text provided")
            
        if initialized_gemini_model is None:
            raise HTTPException(status_code=500, detail="Gemini model not initialized")

        # Generate response using Gemini
        response = initialized_gemini_model.generate_content(
            user_text,
            generation_config={
                "max_output_tokens": 1024,
                "temperature": 0.7,
                "top_p": 0.8,
                "top_k": 40
            }
        )

        # Process the response
        if not response.text:
            raise HTTPException(status_code=500, detail="Empty response from Gemini")

        # Clean the response for speech
        spoken_response = clean_text_for_speech(response.text)
        
        return {
            "spoken_response": spoken_response,
            "raw_response": response.text,
            "chat_history": chat_history + [
                {"role": "user", "content": user_text},
                {"role": "assistant", "content": response.text}
            ]
        }
    except Exception as e:
        logger.error(f"Error in voice query: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/start-voice-session")
async def start_voice_session():
    try:
        return {
            "status": "success",
            "message": "Voice session started",
            "spoken_response": "Hello! I'm your voice assistant. How can I help you today?"
        }
    except Exception as e:
        logger.error(f"Error starting voice session: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# ...existing code...

@app.post("/fetch-youtube-videos") # New endpoint for fetching YouTube videos
async def fetch_youtube_videos_endpoint(request: YouTubeSearchRequest):
    logger.info(f"Received request to fetch YouTube videos for query: {request.query}")
    if not YOUTUBE_API_KEY: # Make sure YOUTUBE_API_KEY is loaded globally in main.py
        logger.error("YOUTUBE_API_KEY is not configured on the server.")
        raise HTTPException(status_code=500, detail="YouTube API key is not configured on the server.")
    try:
        video_data = search_youtube_videos(request.query, max_results=4) # Ensure search_youtube_videos is imported
        if video_data.get("error"):
            logger.error(f"Error from search_youtube_videos: {video_data.get('error')}")
            raise HTTPException(status_code=500, detail=video_data.get("error"))
        return JSONResponse(status_code=200, content=video_data)
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error fetching YouTube videos for query '{request.query}': {str(e)}", exc_info=True)
        return JSONResponse(
            status_code=500,
            content={"error": f"An unexpected error occurred while fetching videos: {str(e)}"}
        )
    
@app.post("/summarize-voice-chat") # New endpoint for summarizing voice chat
async def summarize_voice_chat_endpoint(request: SummarizeChatRequest):
    logger.info(f"Received request to summarize chat history of length: {len(request.chat_history)}")
    try:
        summary_data = summarize_conversation_with_gemini(request.chat_history)
        return JSONResponse(status_code=200, content=summary_data)
    except Exception as e:
        logger.error(f"Error summarizing chat: {str(e)}", exc_info=True)
        return JSONResponse(
            status_code=500,
            content={"summary_text": f"An error occurred during summarization: {str(e)}"}
        )

@app.get("/get-saved-chats")
async def get_saved_chats():
    try:
        if supabase is None:
            logger.error("Supabase client not initialized. Cannot fetch saved chats.")
            raise HTTPException(status_code=500, detail="Database client not initialized.")

        response = supabase.table("chat_history").select("id, filename, chat_data").execute()

        if hasattr(response, 'error') and response.error:
            logger.error(f"Error fetching saved chats: {response.error}")
            raise HTTPException(status_code=500, detail=f"Error fetching saved chats: {response.error}")

        fetched_chats = []
        if response.data is not None:
            for chat_entry in response.data:
                processed_entry = {"id": chat_entry.get("id"), "filename": chat_entry.get("filename")}
                raw_chat_data = chat_entry.get('chat_data')
                if isinstance(raw_chat_data, str):
                    try:
                        processed_entry['chat_history'] = json.loads(raw_chat_data)
                    except json.JSONDecodeError:
                        logger.error(f"Failed to parse chat_data JSON for chat ID {chat_entry.get('id', 'unknown')}. Data: {raw_chat_data[:100]}...")
                        processed_entry['chat_history'] = []
                elif isinstance(raw_chat_data, (list, dict)):
                    processed_entry['chat_history'] = raw_chat_data
                elif raw_chat_data is None:
                    processed_entry['chat_history'] = []
                else:
                    logger.warning(f"Unexpected type for chat_data for chat ID {chat_entry.get('id', 'unknown')}: {type(raw_chat_data)}. Data: {raw_chat_data}")
                    processed_entry['chat_history'] = []
                fetched_chats.append(processed_entry)

        logger.info(f"Successfully fetched and parsed {len(fetched_chats)} saved chats.")
        return JSONResponse(status_code=200, content=fetched_chats)

    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error fetching saved chats: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error fetching saved chats: {str(e)}")

@app.post("/delete-chat")
async def delete_chat(request: DeleteChatRequest):
    try:
        if supabase is None:
            logger.error("Supabase client not initialized. Cannot delete chat.")
            raise HTTPException(status_code=500, detail="Database client not initialized.")

        response = supabase.table("chat_history").delete().eq("id", request.id).execute()
        logger.info(f"Delete operation response for id {request.id}: {response}")

        if hasattr(response, 'error') and response.error:
            logger.error(f"Error in DELETE operation for id {request.id}: {response.error}")
            raise HTTPException(status_code=500, detail=f"Database error during delete: {response.error}")

        deleted_count = getattr(response, 'count', 0)
        if deleted_count > 0:
            logger.info(f"Successfully deleted chat with id: {request.id}")
            return JSONResponse(status_code=200, content={"message": f"Chat with id {request.id} deleted successfully"})
        else:
            logger.warning(f"Delete operation for id {request.id} reported no rows affected. Chat might not exist.")
            raise HTTPException(status_code=404, detail=f"Chat with id {request.id} not found or already deleted")

    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error deleting chat with id {request.id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting chat: {str(e)}")

@app.post("/update-chat-filename")
async def update_chat_filename(request: UpdateChatFilenameRequest):
    try:
        if supabase is None:
            logger.error("Supabase client not initialized. Cannot update chat filename.")
            raise HTTPException(status_code=500, detail="Database client not initialized.")

        existing_chat_response = supabase.table("chat_history").select("id").eq("id", request.id).execute()
        logger.info(f"Check chat existence response for update id {request.id}: {existing_chat_response}")

        if hasattr(existing_chat_response, 'error') and existing_chat_response.error:
            logger.error(f"Error checking existence for id {request.id} during update: {existing_chat_response.error}")
            raise HTTPException(status_code=500, detail=f"Database error during existence check: {existing_chat_response.error}")

        if not existing_chat_response.data:
            logger.error(f"Chat with id {request.id} not found in database for update.")
            raise HTTPException(status_code=404, detail=f"Chat with id {request.id} not found")

        filename_check_response = supabase.table("chat_history").select("id").eq("filename", request.filename).execute()
        logger.info(f"Check filename existence response for filename '{request.filename}' during update: {filename_check_response}")

        if hasattr(filename_check_response, 'error') and filename_check_response.error:
            logger.error(f"Error checking filename existence for '{request.filename}': {filename_check_response.error}")
            raise HTTPException(status_code=500, detail=f"Database error during filename check: {filename_check_response.error}")

        if filename_check_response.data and len(filename_check_response.data) > 0:
            if filename_check_response.data[0]["id"] != request.id:
                logger.error(f"Filename '{request.filename}' already exists for another chat (ID: {filename_check_response.data[0]['id']}).")
                raise HTTPException(status_code=400, detail=f"Filename '{request.filename}' is already in use by another chat")

        updated_data = {"filename": request.filename}
        response = supabase.table("chat_history").update(updated_data).eq("id", request.id).execute()
        logger.info(f"Update filename operation response for id {request.id}: {response}")

        if hasattr(response, 'error') and response.error:
            logger.error(f"Error in UPDATE operation for id {request.id}: {response.error}")
            raise HTTPException(status_code=500, detail=f"Database error during update: {response.error}")

        updated_count = getattr(response, 'count', len(response.data) if response.data is not None else 0)
        if updated_count > 0:
            logger.info(f"Successfully updated filename for chat with id: {request.id} to {request.filename}")
            return JSONResponse(status_code=200, content={"message": f"Chat filename updated to '{request.filename}'"})
        else:
            logger.warning(f"Update filename operation for id {request.id} reported no rows affected. Filename might be the same as current.")
            return JSONResponse(status_code=200, content={"message": f"Chat filename confirmed as '{request.filename}' (no change needed)"})

    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error updating chat filename with id {request.id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error updating chat filename: {str(e)}")

from contextlib import asynccontextmanager

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup: Nothing to initialize for Gemini client
    yield
    # Shutdown
    try:
        # No need to close the Gemini client as it doesn't require explicit closing
        logger.info("Shutdown event: Gemini client doesn't require explicit closing")
    except Exception as e:
        logger.error(f"Error during shutdown: {e}")
    
    # Clean up Supabase client if needed
    logger.info("Application shutting down. Cleaning up resources.")

# Serve static files (game thumbnails, etc.)
app.mount("/static", StaticFiles(directory=str(Path(__file__).parent.parent / "public")), name="static")

# Serve game files statically (only for HTML/JS games like ak02)
# Keep the mount for ak02 HTML game
app.mount("/games/ak02", StaticFiles(directory=str(Path(__file__).parent.parent / "backend/ak02/Gun-Shoot-Game-using-HTML-javaScript-main"), html=True), name="ak02")

# Remove static mounts for Python games
# app.mount("/games/ak01", StaticFiles(directory=str(games_directory / "backend/ak01/snake-game-main"), html=True), name="ak01") # Removed
# app.mount("/games/speed-racer", StaticFiles(directory=str(games_directory / "backend/SPEED-RACER-master/SPEED-RACER-master"), html=True), name="speed-racer") # Removed

@app.post("/generate-quiz", response_model=QuizResponse)
async def generate_quiz(request: QuizRequest):
    """
    Generate a quiz for a specific topic using the AI model.
    The questions are personalized based on the user's learning level.
    """
    logger.info(f"Received quiz request for topic: {request.topic}, difficulty: {request.difficulty}")
    
    try:
        # Set up quiz generation prompt based on difficulty and question type
        if request.question_type == "multiple_choice":
            if request.difficulty == "easy":
                points_per_question = 5
                system_prompt = f"Create {request.num_questions} beginner-friendly multiple-choice questions about {request.topic}."
            elif request.difficulty == "hard":
                points_per_question = 15
                system_prompt = f"Create {request.num_questions} challenging multiple-choice questions about {request.topic} that test deep understanding."
            else:  # medium (default)
                points_per_question = 10
                system_prompt = f"Create {request.num_questions} intermediate-level multiple-choice questions about {request.topic}."
        elif request.question_type == "true_false":
            points_per_question = 5
            system_prompt = f"Create {request.num_questions} true/false questions about {request.topic}."
        else:  # open_ended
            points_per_question = 20
            system_prompt = f"Create {request.num_questions} open-ended questions about {request.topic}."
        
        user_prompt = f"""
        Generate a quiz with {request.num_questions} questions about {request.topic}.
        
        Each question must include:
        1. The question text
        2. Answer options (for multiple choice) with exactly one correct answer
        3. The correct answer
        4. A brief explanation of why that answer is correct
        
        Format your response as JSON with this structure:
        {{
          "questions": [
            {{
              "question": "Question text here?",
              "options": ["Option A", "Option B", "Option C", "Option D"],
              "correct_answer": "Option A",
              "explanation": "Explanation of why Option A is correct"
            }}
          ]
        }}
        
        Make sure the quiz is appropriate for {request.difficulty} difficulty level.
        """
        
        # Get personalization data for better quiz adaptation
        try:
            # Create a personalization agent for a generic user
            from agents.personalization.agent import PersonalizationAgent
            personalization_agent = PersonalizationAgent("quiz_user")
            
            # Get personalization data for this topic
            personalization_data = personalization_agent.process_query(request.topic)
            
            # Add personalization context to the prompt if available
            if personalization_data and "level" in personalization_data:
                user_prompt += f"\nNote: Adjust questions for a {personalization_data['level']} learner."
        except Exception as e:
            logger.warning(f"Could not use personalization agent: {e}")
        
        # Generate quiz using Gemini
        if initialized_gemini_model:
            model = initialized_gemini_model
            
            # Send the prompts to Gemini
            response = model.generate_content(
                f"{system_prompt}\n\n{user_prompt}",
                generation_config={
                    "max_output_tokens": 4096,
                    "temperature": 0.7,
                    "top_p": 0.8
                }
            )
            
            # Parse the response to extract the questions
            response_text = response.text
            
            # Extract JSON from the response (it may be wrapped in Markdown code blocks)
            import json
            import re
            
            # Extract JSON using regex
            json_pattern = re.compile(r'```(?:json)?\s*({[\s\S]*?})```|({[\s\S]*})')
            json_match = json_pattern.search(response_text)
            
            if json_match:
                json_str = json_match.group(1) or json_match.group(2)
                questions_data = json.loads(json_str)
            else:
                # Try parsing the whole response as JSON if no code block is found
                try:
                    questions_data = json.loads(response_text)
                except:
                    # Fallback for when the response isn't valid JSON
                    logger.error(f"Could not extract JSON from model response: {response_text}")
                    questions_data = {"questions": []}
            
            # Format the questions
            formatted_questions = []
            for q in questions_data.get("questions", []):
                # Ensure all required fields are present
                if "question" not in q or "correct_answer" not in q:
                    continue
                
                # For multiple choice, ensure options are provided
                if request.question_type == "multiple_choice" and ("options" not in q or len(q["options"]) < 2):
                    continue
                
                # For true/false, set options to ["True", "False"]
                if request.question_type == "true_false":
                    q["options"] = ["True", "False"]
                
                # Add explanation if missing
                if "explanation" not in q:
                    q["explanation"] = "The correct answer is " + q["correct_answer"]
                
                # Add points
                q["points"] = points_per_question                
                formatted_questions.append(QuizQuestion(**q))
            
            # Calculate total points
            total_points = sum(q.points for q in formatted_questions)
            
            # Return the quiz
            quiz_response = QuizResponse(
                topic=request.topic,
                difficulty=request.difficulty,
                questions=formatted_questions[:request.num_questions],  # Limit to requested number
                total_points=total_points
            )
            
            return quiz_response
        
        else:
            logger.error("Gemini client not initialized. Cannot generate quiz.")
            # Return some default questions as a fallback
            fallback_questions = generate_fallback_questions(request.topic, request.num_questions, request.question_type)
            
            return QuizResponse(
                topic=request.topic,
                difficulty=request.difficulty,
                questions=fallback_questions,
                total_points=sum(q.points for q in fallback_questions)
            )
            
    except Exception as e:
        logger.error(f"Error generating quiz: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate quiz: {str(e)}")

def generate_fallback_questions(topic, num_questions, question_type):
    """Generate fallback questions when the AI model fails."""
    fallback_questions = [
        QuizQuestion(
            question=f"What is one key concept related to {topic}?",
            options=["First concept", "Second concept", "Third concept", "Fourth concept"],
            correct_answer="First concept",
            explanation="This is a default question when AI generation fails.",
            points=10
        )
    ]
    
    for i in range(1, min(num_questions, 5)):
        fallback_questions.append(
            QuizQuestion(
                question=f"Question {i+1} about {topic}?",
                options=[f"Option A for Q{i+1}", f"Option B for Q{i+1}", f"Option C for Q{i+1}", f"Option D for Q{i+1}"],
                correct_answer=f"Option A for Q{i+1}",
                explanation=f"This is fallback question {i+1}.",
                points=10
            )
        )
    
    return fallback_questions

# Endpoint to launch AK01 Python game
@app.get("/launch-ak01-game")
async def launch_ak01_game():
    game_id = "ak01"
    game_script_path = Path(__file__).parent.parent / "backend/ak01/snake-game-main/quiz_snake_game.py"
    game_directory = game_script_path.parent # Get the directory containing the script
    logger.info(f"Attempting to launch AK01 game script: {game_script_path}")
    if not game_script_path.exists():
        logger.error(f"AK01 game script not found at: {game_script_path}")
        raise HTTPException(status_code=404, detail="AK01 game script not found.")

    try:
        # Use subprocess to run the Python script with the correct working directory and pass game_id
        subprocess.Popen([sys.executable, str(game_script_path), game_id], cwd=str(game_directory))
        logger.info("AK01 game script launched successfully on the server.")
        return JSONResponse(status_code=200, content={
            "message": "AK01 game launched on the server. Check the server's display.",
            "note": "The graphical game runs on the server, not directly in your web browser."
        })
    except Exception as e:
        logger.error(f"Error launching AK01 game script: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to launch AK01 game: {e}")

# Endpoint to launch Speed Racer Python game
@app.get("/launch-speed-racer-game")
async def launch_speed_racer_game():
    game_id = "speed-racer"
    game_script_path = Path(__file__).parent.parent / "backend/SPEED-RACER-master/SPEED-RACER-master/Speed Racer.py"
    game_directory = game_script_path.parent # Get the directory containing the script
    logger.info(f"Attempting to launch Speed Racer game script: {game_script_path}")
    if not game_script_path.exists():
        logger.error(f"Speed Racer game script not found at: {game_script_path}")
        raise HTTPException(status_code=404, detail="Speed Racer game script not found.")

    try:
        # Use subprocess to run the Python script with the correct working directory and pass game_id
        subprocess.Popen([sys.executable, str(game_script_path), game_id], cwd=str(game_directory))
        logger.info("Speed Racer game script launched successfully on the server.")
        return JSONResponse(status_code=200, content={
            "message": "Speed Racer game launched on the server. Check the server's display.",
            "note": "The graphical game runs on the server, not directly in your web browser."
        })
    except Exception as e:
        logger.error(f"Error launching Speed Racer game script: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to launch Speed Racer game: {e}")

@app.post("/generate-game-quiz/{game_id}")
async def generate_game_quiz(game_id: str, request: QuizRequest):
    """
    Generate a quiz for a specific game using the AI model and store it temporarily.
    """
    logger.info(f"Received game quiz request for game: {game_id}, topic: {request.topic}, difficulty: {request.difficulty}")

    try:
        # Use the existing quiz generation logic
        quiz_response = await generate_quiz(request) # Re-use the core logic

        # Store the generated quiz data with a unique ID (using game_id for simplicity)
        active_game_quizzes[game_id] = quiz_response.model_dump()
        logger.info(f"Generated and stored quiz for game: {game_id}")

        return JSONResponse(status_code=200, content={
            "message": "Quiz generated and stored",
            "game_id": game_id # Return the identifier used to store the quiz
        })

    except Exception as e:
        logger.error(f"Error generating game quiz for {game_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate game quiz: {str(e)}")

@app.get("/get-game-quiz/{game_id}")
async def get_game_quiz(game_id: str):
    """
    Retrieve the temporarily stored quiz data for a specific game.
    """
    logger.info(f"Received request to get game quiz for game: {game_id}")
    quiz_data = active_game_quizzes.get(game_id)
    if quiz_data:
        logger.info(f"Found quiz data for game: {game_id}")
        # Optionally, remove the quiz after retrieval if it's meant for one-time use
        # del active_game_quizzes[game_id]
        return JSONResponse(status_code=200, content=quiz_data)
    else:
        logger.warning(f"No quiz data found for game: {game_id}")
        raise HTTPException(status_code=404, detail="No quiz data found for this game ID.")

@app.get("/games/memory-match/")
async def serve_memory_match():
    return FileResponse(Path(__file__).parent / "games/memory_match/index.html")

@app.get("/get-game-quiz/memory-match")
async def get_memory_match_quiz():
    """
    Generate quiz questions for the memory match game.
    """
    questions = [
        {"question": "What is 2 + 2?", "answer": "4"},
        {"question": "What is 5 x 5?", "answer": "25"},
        {"question": "What is 10 ÷ 2?", "answer": "5"},
        {"question": "What is 3 x 4?", "answer": "12"},
        {"question": "What is 15 - 7?", "answer": "8"},
        {"question": "What is 6 + 6?", "answer": "12"}
    ]
    return {"questions": questions}

async def summarize_conversation_with_gemini(chat_history):
    try:
        # Convert chat history to a format suitable for summarization
        conversation_text = "\n".join([f"{msg['role']}: {msg['parts'][0]}" for msg in chat_history])
        
        # Use generate_gemini_response to create a summary
        summary = await generate_gemini_response(
            f"Please summarize this conversation briefly:\n\n{conversation_text}",
            []  # Empty chat history since we're just summarizing
        )
        
        return {"summary": summary}
    except Exception as e:
        return {"summary": "Unable to generate summary", "error": str(e)}

@app.post("/api/generate-code", response_model=CodeGenerationResponse)
async def generate_code(request: CodeGenerationRequest):
    """
    Generate HTML/CSS/JavaScript code based on user prompt using AI.
    """
    try:
        logger.info(f"Code generation request received: {request.prompt[:100]}...")
        
        # System prompt for code generation
        system_prompt = """You are an expert web developer and designer. Your task is to generate complete, functional HTML files with embedded CSS and JavaScript based on user prompts.

IMPORTANT GUIDELINES:
1. Always return a complete, self-contained HTML file that can be opened directly in a browser
2. Include all CSS in a <style> tag in the <head> section
3. Include all JavaScript in a <script> tag before the closing </body> tag
4. Use modern, responsive design principles
5. Ensure the code is clean, well-commented, and follows best practices
6. Make the design visually appealing with modern UI/UX patterns
7. Use semantic HTML elements when appropriate
8. Include proper meta tags for responsive design
9. Test that all interactive elements work properly
10. If the prompt is vague, create something creative and impressive

STRUCTURE YOUR RESPONSE:
- Return ONLY the HTML code, no explanations or markdown formatting
- Start with <!DOCTYPE html> and end with </html>
- Do not include any text before or after the HTML code
- Do not wrap the code in backticks or code blocks

USER PROMPT: {prompt}

Generate a complete, functional HTML file:"""

        # Format the system prompt with the user's request
        full_prompt = system_prompt.format(prompt=request.prompt)
        
        # Generate code using Gemini
        if initialized_gemini_model:
            response = await initialized_gemini_model.generate_content_async(full_prompt)
            generated_code = response.text
            
            # Clean up the response to ensure it's just the HTML code
            generated_code = generated_code.strip()

            
            # Remove any markdown formatting if present
            if generated_code.startswith("```html"):
                generated_code = generated_code[7:]
            if generated_code.startswith("```"):
                generated_code = generated_code[3:]
            if generated_code.endswith("```"):
                generated_code = generated_code[:-3]
            
            generated_code = generated_code.strip()
            
            # Ensure the code starts with DOCTYPE
            if not generated_code.startswith("<!DOCTYPE"):
                               generated_code = "<!DOCTYPE html>\n" + generated_code
            
            logger.info("Code generation completed successfully")
            
            return {
                "success": True,
                "code": generated_code,
                "provider": request.provider
            }
        else:
            logger.error("Gemini model not initialized")
            raise HTTPException(status_code=500, detail="AI model not available")
            
    except Exception as e:
        logger.error(f"Error generating code: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "code": """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Error</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            display: flex;
            justify-content: center;
            align-items: center;
            min-height: 100vh;
            margin: 0;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
        }
        .error-container {
            text-align: center;
            padding: 2rem;
            background: rgba(255, 255, 255, 0.1);
            border-radius: 10px;
            backdrop-filter: blur(10px);
        }
    </style>
</head>
<body>
    <div class="error-container">
        <h1>Code Generation Error</h1>
        <p>Sorry, there was an error generating your code. Please try again with a different prompt.</p>
    </div>
</body>
</html>"""
        }

# Enhanced Learning and Analysis Endpoints

@app.post("/api/analyze-code", response_model=CodeAnalysisResponse)
async def analyze_code_comprehensive(request: CodeAnalysisRequest):
    """
    Comprehensive code analysis with educational feedback
    """
    try:
        logger.info(f"Analyzing code for user: {request.user_id}, language: {request.language}")
        
        # Perform comprehensive analysis
        analysis = advanced_analyzer.analyze_code_comprehensive(
            code=request.code,
            language=request.language,
            user_level=request.user_level
        )
        
        # Get personalized feedback if user_id provided
        personalized_feedback = {}
        if request.user_id:
            feedback = learning_agent.get_personalized_feedback(
                user_id=request.user_id,
                code=request.code,
                language=request.language
            )
            personalized_feedback = {
                "user_level": feedback.user_level.value,
                "explanation_complexity": feedback.explanation_complexity.value,
                "suggestions": feedback.suggestions,
                "learning_resources": feedback.learning_resources,
                "next_challenges": feedback.next_challenges
            }
            
            # Update user's learning progress
            learning_agent.update_user_learning_progress(request.user_id, {
                "type": "code_analysis",
                "topic": request.language,
                "complexity": request.user_level,
                "feedback_provided": True
            })
        
        # Convert issues to dictionaries
        issues_dict = []
        for issue in analysis.issues:
            issues_dict.append({
                "line_number": issue.line_number,
                "column": issue.column,
                "issue_type": issue.issue_type.value,
                "severity": issue.severity.value,
                "message": issue.message,
                "suggestion": issue.suggestion,
                "educational_note": issue.educational_note,
                "example_fix": issue.example_fix
            })
        
        response = CodeAnalysisResponse(
            overall_score=analysis.overall_score,
            issues=issues_dict,
            security_score=analysis.security_analysis.security_score,
            performance_score=analysis.performance_analysis.performance_score,
            educational_insights=analysis.educational_insights,
            improvement_roadmap=analysis.improvement_roadmap,
            personalized_feedback=personalized_feedback
        )
        
        logger.info(f"Code analysis completed. Overall score: {analysis.overall_score}")
        return response
        
    except Exception as e:
        logger.error(f"Error in code analysis: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Code analysis failed: {str(e)}")

# Helper function to safely parse JSON strings
def parse_json_field(data, key, default_value):
    value = data.get(key)
    if isinstance(value, str):
        try:
            return json.loads(value)
        except json.JSONDecodeError:
            logger.warning(f"Failed to decode JSON for key {key}: {value}")
            return default_value
    return value if value is not None else default_value

@app.get("/api/user-profile/{user_id}", response_model=UserLearningProfile)
async def get_user_learning_profile(user_id: str):
    """
    Get user's learning profile and progress
    """
    try:
        logger.info(f"Fetching learning profile for user: {user_id}")
        
        profile = learning_agent.get_user_profile(user_id)
        
        # Parse JSON fields that might be stored as strings
        profile['learning_preferences'] = parse_json_field(profile, 'learning_preferences', {})
        profile['goals'] = parse_json_field(profile, 'goals', [])
        profile['weak_topics'] = parse_json_field(profile, 'weak_topics', [])
        profile['interaction_types'] = parse_json_field(profile, 'interaction_types', {})
        profile['topic_progress'] = parse_json_field(profile, 'topic_progress', {})
        profile['achievements'] = parse_json_field(profile, 'achievements', [])

        # Generate current recommendations
        recommendations = []
        skill_level = SkillLevel(profile.get('skill_level', 'beginner'))
        
        if skill_level == SkillLevel.BEGINNER:
            recommendations = [
                "Practice writing simple functions",
                "Learn about variables and data types",
                "Understand basic control structures (if/else, loops)"
            ]
        elif skill_level == SkillLevel.INTERMEDIATE:
            recommendations = [
                "Explore object-oriented programming concepts",
                "Learn about APIs and asynchronous programming",
                "Practice debugging techniques"
            ]
        else:
            recommendations = [
                "Study advanced design patterns",
                "Contribute to open-source projects",
                "Mentor junior developers"
            ]
        
        return UserLearningProfile(
            user_id=user_id,
            skill_level=profile.get('skill_level', 'beginner'),
            knowledge_areas=profile.get('knowledge_areas', {}),
            learning_progress=profile.get('learning_progress', {}),
            recommendations=recommendations,
            learning_preferences=profile.get('learning_preferences', {}),
            goals=profile.get('goals', []),
            weak_topics=profile.get('weak_topics', []),
            interaction_types=profile.get('interaction_types', {}),
            topic_progress=profile.get('topic_progress', {}),
            achievements=profile.get('achievements', [])
        )
        
    except Exception as e:
        logger.error(f"Error fetching user profile: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch user profile: {str(e)}")

@app.get("/api/user-activity/{user_id}")
async def get_user_activity(user_id: str):
    """
    Get user's learning activity data for the dashboard
    """
    try:
        logger.info(f"Fetching user activity for: {user_id}")
        
        # Load user profile to get activity data
        profile = learning_agent.get_user_profile(user_id)
        
        # Create realistic activity data based on user profile
        from datetime import datetime, timedelta
        now = datetime.now()
        
        # Generate activity data for the last 7 days
        daily_activity = []
        for i in range(7):
            date = (now - timedelta(days=i)).strftime('%Y-%m-%d')
            # Generate realistic activity based on user's learning progress
            session_count = max(0, len(profile.get('learning_interactions', [])) - i * 2)
            time_spent = max(0, session_count * 15)  # 15 minutes per session
            concepts_learned = max(0, session_count // 2)
            
            daily_activity.append({
                "date": date,
                "sessions": session_count,
                "timeSpent": time_spent,
                "conceptsLearned": concepts_learned
            })
        
        # Calculate skill sessions
        skill_sessions = {}
        for area, progress in profile.get('knowledge_areas', {}).items():
            skill_sessions[area] = max(1, int(progress * 10))  # Convert progress to session count
        
        # Calculate metrics
        total_interactions = len(profile.get('learning_interactions', []))
        total_time = sum(activity['timeSpent'] for activity in daily_activity)
        concepts_learned = len(profile.get('learned_concepts', []))
        
        # Calculate current streak
        current_streak = 0
        for activity in daily_activity:
            if activity['sessions'] > 0:
                current_streak += 1
            else:
                break
        
        # Get recent topics
        recent_topics = list(profile.get('knowledge_areas', {}).keys())[:3]
        
        activity_data = {
            "sessions": total_interactions,
            "totalTime": total_time,
            "lastActive": now.isoformat(),
            "currentStreak": current_streak,
            "dailyActivity": daily_activity,
            "skillSessions": skill_sessions,
            "recentTopics": recent_topics
        }
        
        return activity_data
        
    except Exception as e:
        logger.error(f"Error fetching user activity: {str(e)}")
        # Return empty activity data for new users
        return {
            "sessions": 0,
            "totalTime": 0,
            "lastActive": datetime.now().isoformat(),
            "currentStreak": 0,
            "dailyActivity": [],
            "skillSessions": {},
            "recentTopics": []
        }

@app.post("/api/assess-skill-level")
async def assess_user_skill_level(request: SkillAssessmentRequest):
    """
    Assess user's skill level based on code samples
    """
    try:
        logger.info(f"Assessing skill level for user: {request.user_id}")
        
        # Assess skill level
        assessed_level = learning_agent.assess_skill_level(request.user_id, request.code_samples)
        
        # Update user profile with assessed skill level
        profile = learning_agent.get_user_profile(request.user_id)
        profile['skill_level'] = assessed_level.value
        profile['updated_at'] = datetime.now().isoformat()
        learning_agent._save_user_profile(request.user_id, profile)
        
        return {
            "user_id": request.user_id,
            "assessed_skill_level": assessed_level.value,
            "confidence": 0.85,  # Placeholder confidence score
            "recommendations": learning_agent._recommend_learning_resources(assessed_level, request.language)
        }
        
    except Exception as e:
        logger.error(f"Error in skill assessment: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Skill assessment failed: {str(e)}")

@app.post("/api/learning-interaction")
async def record_learning_interaction(request: LearningInteractionRequest):
    """
    Record a learning interaction and update user progress
    """
    try:
        logger.info(f"Recording learning interaction for user: {request.user_id}")
        
        interaction_data = {
            "type": request.interaction_type,
            "topic": request.topic,
            "complexity": request.complexity,
            "feedback_provided": True,
            "timestamp": datetime.now().isoformat()
        }
        
        # Update user's learning progress
        updated_profile = learning_agent.update_user_learning_progress(
            request.user_id, 
            interaction_data
        )
        
        return {
            "success": True,
            "message": "Learning interaction recorded successfully",
            "updated_interactions_count": updated_profile['interactions_count'],
            "current_skill_level": updated_profile['skill_level']
        }
        
    except Exception as e:
        logger.error(f"Error recording learning interaction: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to record interaction: {str(e)}")

@app.post("/api/analyze-coding-style")
async def analyze_user_coding_style(request: dict):
    """
    Analyze user's coding style from code samples for personalization
    """
    try:
        user_id = request.get('user_id', 'anonymous')
        code_samples = request.get('code_samples', [])
        
        logger.info(f"Analyzing coding style for user: {user_id}")
        
        # Analyze coding style
        coding_style = advanced_personalization_agent.analyze_coding_style(user_id, code_samples)
        
        # Update user profile with coding style
        profile = advanced_personalization_agent.get_advanced_user_profile(user_id)
        profile['coding_style'] = {
            "preferred_naming": coding_style.preferred_naming,
            "indentation": coding_style.indentation,
            "comment_frequency": coding_style.comment_frequency,
            "function_length_preference": coding_style.function_length_preference,
            "error_handling_style": coding_style.error_handling_style,
            "preferred_paradigm": coding_style.preferred_paradigm.value
        }
        advanced_personalization_agent._save_advanced_profile(user_id, profile)
        
        return {
            "user_id": user_id,
            "coding_style": profile['coding_style'],
            "analysis_summary": "Coding style preferences have been analyzed and saved",
            "personalization_enabled": True
        }
        
    except Exception as e:
        logger.error(f"Error analyzing coding style: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Coding style analysis failed: {str(e)}")

@app.post("/api/update-project-context")
async def update_project_context(request: dict):
    """
    Update user's project context for better personalization
    """
    try:
        user_id = request.get('user_id', 'anonymous')
        project_data = request.get('project_data', {})
        
        logger.info(f"Updating project context for user: {user_id}")
        
        # Update project context
        advanced_personalization_agent.update_project_context(user_id, project_data)
        
        return {
            "success": True,
            "message": "Project context updated successfully",
            "project_id": project_data.get('project_id', 'unknown')
        }
        
    except Exception as e:
        logger.error(f"Error updating project context: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Project context update failed: {str(e)}")

@app.post("/api/get-personalized-suggestions")
async def get_personalized_suggestions(request: dict):
    """
    Get personalized suggestions based on user profile and current code
    """
    try:
        user_id = request.get('user_id', 'anonymous')
        current_code = request.get('code', '')
        task_type = request.get('task_type', 'general')
        
        logger.info(f"Generating personalized suggestions for user: {user_id}")
        
        # Get personalized suggestions
        suggestions = advanced_personalization_agent.get_personalized_suggestions(
            user_id, current_code, task_type
        )
        
        # Get learning insights
        insights = advanced_personalization_agent.generate_learning_insights(user_id)
        
        return {
            "personalized_suggestions": suggestions,
            "learning_insights": [
                {
                    "type": insight.insight_type,
                    "message": insight.message,
                    "confidence": insight.confidence,
                    "suggestion": insight.actionable_suggestion,
                    "learning_path": insight.learning_path
                } for insight in insights
            ],
            "user_id": user_id
        }
        
    except Exception as e:
        logger.error(f"Error generating personalized suggestions: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Personalized suggestions failed: {str(e)}")

@app.post("/api/refactor-code")
async def get_refactoring_suggestions(request: dict):
    """
    Get intelligent code refactoring suggestions
    """
    try:
        code = request.get('code', '')
        language = request.get('language', 'python')
        user_level = request.get('user_level', 'intermediate')
        
        logger.info(f"Generating refactoring suggestions for {language} code")
        
        # Analyze code quality
        quality_metrics = code_refactoring_agent.analyze_code_quality(code, language)
        
        # Get refactoring suggestions
        refactoring_suggestions = code_refactoring_agent.suggest_refactoring_improvements(
            code, language, user_level
        )
        
        # Get performance optimizations
        performance_optimizations = code_refactoring_agent.suggest_performance_optimizations(
            code, language, user_level
        )
        
        return {
            "quality_metrics": {
                "readability_score": quality_metrics.readability_score,
                "maintainability_score": quality_metrics.maintainability_score,
                "performance_score": quality_metrics.performance_analysis.performance_score,
                "documentation_score": quality_metrics.documentation_score,
                "complexity_score": quality_metrics.complexity_score,
                "security_score": quality_metrics.security_score
            },
            "refactoring_suggestions": [
                {
                    "id": suggestion.suggestion_id,
                    "type": suggestion.type.value,
                    "title": suggestion.title,
                    "description": suggestion.description,
                    "priority": suggestion.priority.value,
                    "difficulty": suggestion.difficulty,
                    "code_before": suggestion.code_before,
                    "code_after": suggestion.code_after,
                    "explanation": suggestion.explanation,
                    "benefits": suggestion.benefits,
                    "learning_opportunity": suggestion.learning_opportunity,
                    "estimated_time": suggestion.estimated_time
                } for suggestion in refactoring_suggestions
            ],
            "performance_optimizations": [
                {
                    "id": opt.optimization_id,
                    "issue": opt.performance_issue,
                    "current_complexity": opt.current_complexity,
                    "optimized_complexity": opt.optimized_complexity,
                    "suggestion": opt.suggestion,
                    "code_example": opt.code_example,
                    "impact": opt.impact_level.value,
                    "explanation": opt.explanation
                } for opt in performance_optimizations
            ]
        }
        
    except Exception as e:
        logger.error(f"Error generating refactoring suggestions: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Refactoring analysis failed: {str(e)}")

@app.post("/api/make-responsive")
async def make_component_responsive(request: dict):
    """
    Get suggestions for making components responsive
    """
    try:
        code = request.get('code', '')
        component_type = request.get('component_type', 'web')
        
        logger.info(f"Generating responsive design suggestions for {component_type} component")
        
        # Get responsive suggestions
        suggestions = code_refactoring_agent.generate_responsive_suggestions(code, component_type)
        
        return {
            "responsive_suggestions": suggestions,
            "component_type": component_type,
            "best_practices": [
                "Use CSS Grid or Flexbox for flexible layouts",
                "Implement mobile-first design approach",
                "Use relative units instead of fixed pixels",
                "Optimize for touch interfaces",
                "Consider performance on mobile devices"
            ]
        }
        
    except Exception as e:
        logger.error(f"Error generating responsive suggestions: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Responsive design suggestions failed: {str(e)}")

@app.post("/api/security-analysis")
async def analyze_security_vulnerabilities(request: dict):
    """
    Detect security vulnerabilities and bugs in code
    """
    try:
        code = request.get('code', '')
        language = request.get('language', 'python')
        
        logger.info(f"Analyzing security vulnerabilities in {language} code")
        
        # Detect security vulnerabilities
        security_issues = security_bug_agent.detect_security_vulnerabilities(code, language)
        
        # Detect potential bugs
        bug_reports = security_bug_agent.detect_bugs(code, language)
        
        # Generate security summary
        security_summary = security_bug_agent.generate_security_summary(security_issues, bug_reports)
        
        return {
            "security_summary": security_summary,
            "security_issues": [
                {
                    "id": issue.issue_id,
                    "type": issue.vulnerability_type.value,
                    "severity": issue.severity.value,
                    "title": issue.title,
                    "description": issue.description,
                    "location": issue.location,
                    "code_snippet": issue.code_snippet,
                    "recommendation": issue.recommendation,
                    "cwe_id": issue.cwe_id,
                    "learning_resource": issue.learning_resource,
                    "fix_example": issue.fix_example
                } for issue in security_issues
            ],
            "bug_reports": [
                {
                    "id": bug.bug_id,
                    "type": bug.bug_type.value,
                    "severity": bug.severity.value,
                    "title": bug.title,
                    "description": bug.description,
                    "location": bug.location,
                    "code_snippet": bug.code_snippet,
                    "fix_suggestion": bug.fix_suggestion,
                    "prevention_tip": bug.prevention_tip,
                    "test_suggestion": bug.test_suggestion
                } for bug in bug_reports
            ]
        }
        
    except Exception as e:
        logger.error(f"Error in security analysis: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Security analysis failed: {str(e)}")

@app.post("/api/explain-algorithm")
async def explain_algorithm_comprehensive(request: dict):
    """
    Provide comprehensive algorithm explanation with visualizations
    """
    try:
        code = request.get('code', '')
        language = request.get('language', 'python')
        user_level = request.get('user_level', 'intermediate')
        
        logger.info(f"Generating comprehensive algorithm explanation for {language} code")
        
        # Generate comprehensive explanation
        explanation = algorithm_explanation_agent.generate_comprehensive_explanation(
            code, language, user_level
        )
        
        # Create visual representation
        visual_representation = algorithm_explanation_agent.create_visual_representation(
            code, explanation.algorithm_type
        )
        
        return {
            "algorithm_explanation": {
                "name": explanation.algorithm_name,
                "type": explanation.algorithm_type.value,
                "description": explanation.description,
                "purpose": explanation.purpose,
                "time_complexity": explanation.time_complexity.value,
                "space_complexity": explanation.space_complexity.value,
                "key_concepts": explanation.key_concepts,
                "learning_objectives": explanation.learning_objectives,
                "common_mistakes": explanation.common_mistakes,
                "optimization_opportunities": explanation.optimization_opportunities,
                "real_world_applications": explanation.real_world_applications,
                "related_algorithms": explanation.related_algorithms
            },
            "code_flow": [
                {
                    "step": step.step_number,
                    "line": step.line_number,
                    "code": step.code_snippet,
                    "explanation": step.explanation,
                    "variables": step.variables_state,
                    "complexity": step.complexity_contribution,
                    "learning_notes": step.learning_notes
                } for step in explanation.code_flow
            ],
            "visual_representation": {
                "diagram_type": visual_representation.diagram_type,
                "elements": visual_representation.elements,
                "connections": visual_representation.connections,
                "annotations": visual_representation.annotations,
                "step_changes": visual_representation.step_by_step_changes
            }
        }
        
    except Exception as e:
        logger.error(f"Error generating algorithm explanation: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Algorithm explanation failed: {str(e)}")

@app.post("/api/adaptive-learning")
async def track_adaptive_learning(request: dict):
    """
    Track learning progress and adapt explanations to user level
    """
    try:
        user_id = request.get('user_id', 'anonymous')
        interaction_data = request.get('interaction_data', {})
        content = request.get('content', '')
        topic = request.get('topic', '')
        
        logger.info(f"Tracking adaptive learning for user: {user_id}")
        
        # Track learning progress
        updated_profile = advanced_personalization_agent.track_learning_progress(
            user_id, interaction_data
        )
        
        # Adapt explanation to user
        adapted_explanation = advanced_personalization_agent.adapt_explanation_to_user(
            user_id, content, topic
        )
        
        return {
            "adapted_explanation": adapted_explanation,
            "learning_progress": {
                "interactions_count": len(updated_profile.get('learning_history', [])),
                "skill_level": updated_profile.get('skill_progression', {}).get('overall_level', 'beginner'),
                "learning_velocity": updated_profile.get('skill_progression', {}).get('learning_velocity', 0.5),
                "knowledge_areas": updated_profile.get('knowledge_areas', {})
            },
            "personalization_active": True
        }
        
    except Exception as e:
        logger.error(f"Error in adaptive learning: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Adaptive learning failed: {str(e)}")

@app.get("/api/learning-visualization/{user_id}")
async def get_learning_visualization(user_id: str):
    """
    Get comprehensive learning visualization data for the user
    """
    try:
        logger.info(f"Generating learning visualization for user: {user_id}")
        
        # Get advanced user profile
        profile = advanced_personalization_agent.get_advanced_user_profile(user_id)
        
        # Generate learning insights
        insights = advanced_personalization_agent.generate_learning_insights(user_id)
        
        # Calculate knowledge progress
        knowledge_areas = profile.get('knowledge_areas', {})
        knowledge_progress = []
        
        for area, data in knowledge_areas.items():
            if isinstance(data, dict):
                knowledge_progress.append({
                    "area": area,
                    "proficiency": data.get('proficiency', 0),
                    "interactions": data.get('interactions', 0),
                    "last_practiced": data.get('last_practiced', '')
                })
        
        # Learning history analysis
        learning_history = profile.get('learning_history', [])
        recent_activity = learning_history[-10:] if len(learning_history) > 10 else learning_history
        
        return {
            "user_id": user_id,
            "skill_progression": profile.get('skill_progression', {}),
            "knowledge_progress": knowledge_progress,
            "learning_insights": [
                {
                    "type": insight.insight_type,
                    "message": insight.message,
                    "confidence": insight.confidence,
                    "suggestion": insight.actionable_suggestion,
                    "learning_path": insight.learning_path
                } for insight in insights
            ],
            "recent_activity": recent_activity,
            "coding_style": profile.get('coding_style', {}),
            "learning_preferences": profile.get('learning_preferences', {}),
            "interaction_patterns": profile.get('interaction_patterns', {}),
            "visualization_data": {
                "progress_over_time": [
                    {"date": session.get('timestamp', ''), "score": session.get('success_rate', 0.5)}
                    for session in recent_activity
                ],
                "skill_radar": knowledge_progress,
                "learning_velocity_trend": profile.get('skill_progression', {}).get('learning_velocity', 0.5)
            }
        }
        
    except Exception as e:
        logger.error(f"Error generating learning visualization: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Learning visualization failed: {str(e)}")

# Phase 1: Activity Tracking Endpoint
@app.post("/api/activity")
async def log_user_activity(request: dict):
    """
    Log user activity for Phase 1 learning analytics
    """
    try:
        # Extract activity data from request
        activity_data = {
            "user_id": request.get("userId", "guest"),
            "activity_type": request.get("activityType", "unknown"),
            "topic": request.get("topic"),
            "details": request.get("details", {}),
            "duration": request.get("duration", 1),
            "engagement_score": request.get("engagementScore", 0.5),
            "page": request.get("page"),
            "session_time": request.get("sessionTime"),
            "timestamp": datetime.now().isoformat()
        }
        
        logger.info(f"Logging activity: {activity_data['activity_type']} for user {activity_data['user_id']}")
        
        # For Phase 1, we'll just log to console and return success
        # In production, this would save to Supabase user_activities table
        
        return {
            "success": True,
            "message": "Activity logged successfully",
            "activity_id": f"activity_{datetime.now().timestamp()}"
        }
        
    except Exception as e:
        logger.error(f"Error logging user activity: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Activity logging failed: {str(e)}")

# Enhanced profile creation endpoint
@app.post("/api/profile/create-detailed")
async def create_detailed_profile(request: Request):
    """Create a detailed user profile with comprehensive information"""
    try:
        body = await request.json()
        email = body.get("email")
        profile_data = body.get("profileData", {})

        logger.info(f"Received detailed profile creation request. Email: {email}, Profile Data Keys: {list(profile_data.keys())}")
        logger.info(f"Full request body: {body}") # Added logging for full request body

        if not email:
            return JSONResponse(
                status_code=400,
                content={"error": "Email is required"}
            )
        
        # Get user from Supabase by email using service role client
        try:
            # Use service role client for admin operations
            admin_client = service_role_supabase if service_role_supabase else supabase
            
            logger.info(f"Looking up user with email: {email}")
            user_response = admin_client.auth.admin.list_users()
            logger.info(f"Supabase list_users() response: {user_response}")
            user_data = None

            # Correctly iterate over user_response if it's a list, or its .users attribute if it has one
            users_list = user_response.users if hasattr(user_response, 'users') else user_response

            if isinstance(users_list, list):
                for user in users_list:
                    logger.info(f"Checking user: {repr(user.email)}, Target email: {repr(email)}") # Added repr for debug
                    if user.email and user.email.lower() == email.lower():
                        user_data = user
                        break

            if not user_data:
                logger.error(f"User not found for email: {email}")
                return JSONResponse(
                    status_code=404,
                    content={"error": "User not found"}
                )
            
            user_id = user_data.id
            logger.info(f"Found user with ID: {user_id}")
            
            # Create comprehensive profile data matching new schema
            comprehensive_profile = {
                "user_id": user_id,  # Add user_id for proper auth integration
                "email": email,
                "name": profile_data.get("name", ""),
                "age": int(profile_data.get("age", 0)) if profile_data.get("age") else None,
                "education_level": profile_data.get("education", ""),
                "occupation": profile_data.get("occupation", ""),
                "preferred_language": profile_data.get("preferredLanguage", "English"),
                
                # Learning Preferences - ensure proper JSON formatting
                "learning_style": json.dumps(profile_data.get("learningStyle", [])) if isinstance(profile_data.get("learningStyle"), list) else json.dumps([profile_data.get("learningStyle", "visual")]),
                "preferred_mode": profile_data.get("preferredMode", "both"),
                "topics_of_interest": json.dumps(profile_data.get("topicsOfInterest", [])),
                "current_goal": profile_data.get("currentGoal", ""),
                "daily_time": profile_data.get("dailyTime", "30min"),
                
                # Prior Knowledge
                "experience_level": profile_data.get("skillLevel", "beginner"),
                "confidence_levels": json.dumps(profile_data.get("confidence", {})),
                "previous_platforms": json.dumps(profile_data.get("previousPlatforms", [])),
                "current_skills": json.dumps(profile_data.get("currentSkills", [])),
                "interests": json.dumps(profile_data.get("interests", [])),
                
                # Intent & Goals
                "primary_reason": profile_data.get("primaryReason", ""),
                "target_deadline": profile_data.get("deadline") if profile_data.get("deadline") else None,
                "want_reminders": profile_data.get("wantReminders", True),
                "reminder_time": profile_data.get("reminderTime", "09:00"),
                "motivation": profile_data.get("motivation", ""),
                "learning_goals": json.dumps(profile_data.get("learningGoals", [])),
                "time_available": profile_data.get("timeAvailable", "1-3"),
                
                # Accessibility
                "text_size": profile_data.get("textSize", "medium"),
                "visual_mode": profile_data.get("visualMode", "dark"),
                "enable_sound": profile_data.get("enableSound", True),
                
                # Learning Analytics
                "total_interactions": 0,
                "learning_streak": 0,
                "skill_level": profile_data.get("skillLevel", "beginner"),
                "preferred_difficulty": profile_data.get("preferredDifficulty", "medium"),
                "last_activity_date": datetime.now().strftime('%Y-%m-%d %H:%M:%S+00'),
                
                # Learning Preferences (for backward compatibility)
                "learning_preferences": json.dumps({
                    "style": profile_data.get("learningStyle", "visual") if isinstance(profile_data.get("learningStyle"), str) else (profile_data.get("learningStyle", ["visual"])[0] if profile_data.get("learningStyle") else "visual"),
                    "pace": "slow" if profile_data.get("dailyTime") == "15min" else ("fast" if profile_data.get("dailyTime") == "1hr" else "normal"),
                    "confidence": 0.5,
                    "mode": profile_data.get("preferredMode", "both"),
                    "textSize": profile_data.get("textSize", "medium"),
                    "visualMode": profile_data.get("visualMode", "dark"),
                    "enableSound": profile_data.get("enableSound", True)
                }),
                
                # Goals and Progress
                "goals": json.dumps([
                    {
                        "id": int(datetime.now().timestamp()),
                        "text": profile_data.get("currentGoal", ""),
                        "created": datetime.now().isoformat(),
                        "completed": False,
                        "progress": 0,
                        "deadline": profile_data.get("deadline") if profile_data.get("deadline") else None,
                        "reason": profile_data.get("primaryReason", "")
                    }
                ] if profile_data.get("currentGoal") else []),
                "progress": json.dumps({}),
                "achievements": json.dumps([]),
                
                # Onboarding Status
                "onboarding_completed": True,
                "onboarding_completed_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S+00'),
                
                # Additional metadata
                "metadata": json.dumps({
                    "onboardingCompletedAt": datetime.now().isoformat(),
                    "profileVersion": "2.0",
                    "dataSource": "onboarding"
                })
            }
            
            # Add detailed logging
            logger.info(f"Attempting to create profile for email: {email}")
            logger.info(f"Profile data keys: {list(profile_data.keys())}")
            logger.info(f"Comprehensive profile created with {len(comprehensive_profile)} fields")
            
            # Insert or update the profile in Supabase using service role client
            try:
                # Use service role client if available, otherwise fall back to regular client
                client_to_use = service_role_supabase if service_role_supabase else supabase
                client_type = "service_role" if service_role_supabase else "anon"
                
                logger.info(f"Using {client_type} client for profile operation")
                
                profile_response = client_to_use.table("user_profiles").upsert(
                    comprehensive_profile,
                    on_conflict="user_id"
                ).execute()
                
                logger.info(f"Supabase response data: {profile_response.data}")
                if profile_response.error:
                    logger.info(f"Supabase response error: {profile_response.error}")

                if profile_response.data and len(profile_response.data) > 0: # Check if data exists and is not empty
                    logger.info(f"Successfully created/updated profile for {email}")
                    return JSONResponse(
                        status_code=200,
                        content={
                            "message": "Profile created successfully",
                            "profile_id": profile_response.data[0].get("id"), # Safely get the ID
                            "email": email
                        }
                    )
                else:
                    logger.info(f"Supabase upsert successful but no data returned (likely an update). Profile for {email}.")
                    return JSONResponse(
                        status_code=200, # Return 200 OK even if no data is returned, as the operation was successful
                        content={
                            "message": "Profile updated successfully (no new data returned)",
                            "profile_id": None, # No new profile ID if it was an update
                            "email": email
                        }
                    )
            except Exception as e:
                logger.error(f"Database error during profile upsert: {e}", exc_info=True)
                return JSONResponse(
                    status_code=500,
                    content={"error": f"Database error: {e}"}
                )
            
        except Exception as supabase_error:
            logger.error(f"Supabase error: {supabase_error}")
            return JSONResponse(
                status_code=500,
                content={"error": f"Database error: {str(supabase_error)}"}
            )
            
    except Exception as e:
        logger.error(f"Error creating detailed profile: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Server error: {str(e)}"}
        )

# Debug endpoint to test profile data
@app.post("/api/profile/debug")
async def debug_profile_data(request: Request):
    """Debug endpoint to see what data we're receiving"""
    try:
        body = await request.json()
        logger.info(f"Debug - Received data: {json.dumps(body, indent=2)}")
        
        return JSONResponse(
            status_code=200,
            content={
                "message": "Debug data received",
                "received_data": body
            }
        )
    except Exception as e:
        logger.error(f"Debug endpoint error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": str(e)}
        )

class ProfileUpdateRequest(BaseModel):
    user_id: str
    updates: Dict[str, Any]

@app.put("/api/profile/update")
async def update_user_profile(request: Request):
    """Update user profile with partial data"""
    try:
        body = await request.json()
        email = body.get("email")
        update_data = body.get("updateData", {})
        
        if not email:
            return JSONResponse(
                status_code=400,
                content={"error": "Email is required"}
            )
        
        # Use service role client for admin operations
        admin_client = service_role_supabase if service_role_supabase else supabase
        
        logger.info(f"Updating profile for email: {email}")
        logger.info(f"Update data: {list(update_data.keys())}")
        
        # Get current user profile
        profile_response = admin_client.table("user_profiles").select("*").eq("email", email).execute()
        
        if not profile_response.data:
            return JSONResponse(
                status_code=404,
                content={"error": "Profile not found"}
            )
        
        current_profile = profile_response.data[0]
        
        # Prepare update data with proper JSON formatting
        formatted_update = {}
        for key, value in update_data.items():
            if key in ['learning_style', 'topics_of_interest', 'current_skills', 'previous_platforms', 'learning_goals', 'interests']:
                # Ensure arrays are properly formatted as JSON
                if isinstance(value, list):
                    formatted_update[key] = json.dumps(value)
                else:
                    formatted_update[key] = value
            elif key in ['confidence_levels', 'learning_preferences', 'metadata']:
                # Ensure objects are properly formatted as JSON
                if isinstance(value, dict):
                    formatted_update[key] = json.dumps(value)
                else:
                    formatted_update[key] = value
            else:
                formatted_update[key] = value
        
        # Add update timestamp
        formatted_update['updated_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S+00')
        
        # Update the profile
        client_to_use = service_role_supabase if service_role_supabase else supabase
        update_response = client_to_use.table("user_profiles").update(formatted_update).eq("email", email).execute()
        
        if update_response.data:
            logger.info(f"Successfully updated profile for {email}")
            return JSONResponse(
                status_code=200,
                content={
                    "message": "Profile updated successfully",
                    "updated_fields": list(formatted_update.keys())
                }
            )
        else:
            logger.error("No data returned from profile update")
            return JSONResponse(
                status_code=500,
                content={"error": "Failed to update profile"}
            )
            
    except Exception as e:
        logger.error(f"Error updating profile: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Server error: {str(e)}"}
        )

@app.get("/api/profile/database/{user_id}")
async def get_user_profile_from_database(user_id: str):
    """
    Get user profile from Supabase database (not file system)
    """
    try:
        logger.info(f"Fetching profile from database for user: {user_id}")
        
        client_to_use = service_role_supabase if service_role_supabase else supabase
        profile_response = None
        
        # Check if user_id looks like an email
        if "@" in user_id:
            logger.info(f"Searching by email: {user_id}")
            profile_response = client_to_use.table("user_profiles").select("*").eq("email", user_id).execute()
        else:
            # Try to search by user_id (UUID)
            logger.info(f"Searching by user_id: {user_id}")
            try:
                profile_response = client_to_use.table("user_profiles").select("*").eq("user_id", user_id).execute()
            except Exception as uuid_error:
                logger.warning(f"Failed to search by user_id (UUID): {uuid_error}")
                profile_response = None
        
        if not profile_response or not profile_response.data:
            logger.warning(f"No profile found for user: {user_id}")
            return JSONResponse(
                status_code=404,
                content={"error": "Profile not found"}
            )
        
        profile_data = profile_response.data[0]
        
        # Parse JSON fields that are stored as strings
        json_fields = [
            'learning_style', 'topics_of_interest', 'current_skills', 'previous_platforms', 
            'learning_goals', 'interests', 'confidence_levels', 'learning_preferences', 
            'metadata', 'goals', 'weak_topics'
        ]
        
        for field in json_fields:
            if field in profile_data and isinstance(profile_data[field], str):
                try:
                    profile_data[field] = json.loads(profile_data[field])
                except json.JSONDecodeError:
                    logger.warning(f"Failed to parse JSON field {field}: {profile_data[field]}")
                    # Set to appropriate default value
                    if field in ['confidence_levels', 'learning_preferences', 'metadata']:
                        profile_data[field] = {}
                    else:
                        profile_data[field] = []
        
        logger.info(f"Successfully retrieved and parsed profile for {user_id}")
        return JSONResponse(
            status_code=200,
            content=profile_data
        )
        
    except Exception as e:
        logger.error(f"Error retrieving profile from database: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Server error: {str(e)}"}
        )
